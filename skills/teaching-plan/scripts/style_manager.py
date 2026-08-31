# -*- coding: utf-8 -*-
"""
StyleManager: Manages document styles, typography rules, XML properties,
and hierarchical style overrides for teaching plan generation.
"""

import json
import os
import copy
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from technical_text import split_technical_text

def _resolve_default_styles_path():
    p1 = os.path.join(os.path.dirname(__file__), "styles", "default_styles.json")
    if os.path.exists(p1):
        return p1
    p2 = os.path.join(os.path.dirname(__file__), "..", "styles", "default_styles.json")
    if os.path.exists(p2):
        return os.path.abspath(p2)
    return p1

DEFAULT_STYLES_PATH = _resolve_default_styles_path()


def contains_cjk(text):
    """Return True when a run contains Chinese/Japanese/Korean characters."""
    return any(
        "\u3400" <= char <= "\u4dbf"
        or "\u4e00" <= char <= "\u9fff"
        or "\uf900" <= char <= "\ufaff"
        for char in str(text)
    )

def deep_merge(target, source):
    """Recursively merges source dict into target dict."""
    if not isinstance(source, dict):
        return source
    result = copy.deepcopy(target)
    for key, value in source.items():
        if key in result and isinstance(result[key], dict) and isinstance(value, dict):
            result[key] = deep_merge(result[key], value)
        else:
            result[key] = copy.deepcopy(value)
    return result

class StyleManager:
    def __init__(self, styles_config=None, style_override=None):
        """
        Initialize StyleManager with base styles and optional overrides.
        :param styles_config: Path to JSON config file or dict of styles. If None, loads default_styles.json.
        :param style_override: Optional dict or JSON path containing style overrides.
        """
        if styles_config is None:
            if os.path.exists(DEFAULT_STYLES_PATH):
                with open(DEFAULT_STYLES_PATH, "r", encoding="utf-8") as f:
                    self.styles = json.load(f)
            else:
                self.styles = {}
        elif isinstance(styles_config, str):
            if os.path.exists(styles_config):
                with open(styles_config, "r", encoding="utf-8") as f:
                    self.styles = json.load(f)
            else:
                # 尝试从 styles 目录检索
                p_local = os.path.join(os.path.dirname(__file__), "styles", f"{styles_config}.json")
                p_parent = os.path.join(os.path.dirname(__file__), "..", "styles", f"{styles_config}.json")
                if os.path.exists(p_local):
                    with open(p_local, "r", encoding="utf-8") as f:
                        self.styles = json.load(f)
                elif os.path.exists(p_parent):
                    with open(p_parent, "r", encoding="utf-8") as f:
                        self.styles = json.load(f)
                else:
                    raise FileNotFoundError(f"未找到样式文件: {styles_config}")
        elif isinstance(styles_config, dict):
            self.styles = copy.deepcopy(styles_config)
        else:
            raise ValueError(f"Unsupported styles_config type: {type(styles_config)}")

        # Apply global style override if provided
        if style_override:
            if isinstance(style_override, str):
                if os.path.exists(style_override):
                    with open(style_override, "r", encoding="utf-8") as f:
                        override_dict = json.load(f)
                else:
                    override_dict = json.loads(style_override)
            elif isinstance(style_override, dict):
                override_dict = style_override
            else:
                raise ValueError(f"Unsupported style_override type: {type(style_override)}")
            self.styles = deep_merge(self.styles, override_dict)

    def get_style(self, *keys, default=None):
        """Get nested style value by keys path."""
        curr = self.styles
        for k in keys:
            if isinstance(curr, dict) and k in curr:
                curr = curr[k]
            else:
                return default
        return curr

    def apply_page_setup(self, section, page_override=None):
        """Configure page margins and dimensions."""
        page_conf = self.get_style("page", default={})
        if page_override:
            page_conf = deep_merge(page_conf, page_override)

        section.page_width = Pt(page_conf.get("page_width_pt", 595.3))
        section.page_height = Pt(page_conf.get("page_height_pt", 841.9))
        section.top_margin = Pt(page_conf.get("top_margin_pt", 56.7))
        section.bottom_margin = Pt(page_conf.get("bottom_margin_pt", 56.7))
        section.left_margin = Pt(page_conf.get("left_margin_pt", 84.6))
        section.right_margin = Pt(page_conf.get("right_margin_pt", 84.6))

    def parse_alignment(self, align_str):
        """Convert string alignment to WD_ALIGN_PARAGRAPH enum."""
        if isinstance(align_str, WD_ALIGN_PARAGRAPH):
            return align_str
        align_map = {
            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
            "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
            "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        return align_map.get(str(align_str).upper(), WD_ALIGN_PARAGRAPH.LEFT)

    def add_run(self, paragraph, text, font_name=None, size_pt=None, bold=False, color_rgb=None, ascii_font=None,
                font_hint=None):
        """Add styled text run with East Asia and Western font XML settings."""
        fonts_conf = self.get_style("fonts", default={})
        substitutions = fonts_conf.get("substitutions", {})
        if not isinstance(substitutions, dict):
            substitutions = {}
        requested_fn = font_name or fonts_conf.get("default_font", "宋体")
        requested_af = ascii_font or fonts_conf.get("ascii_font", "Times New Roman")
        fn = substitutions.get(requested_fn, requested_fn)
        af = substitutions.get(requested_af, requested_af)
        # LibreOffice on some hosts ignores w:eastAsia when w:ascii/w:hAnsi
        # points at a Western-only font. Use the CJK font for every font slot in
        # a run that actually contains CJK glyphs; Latin-only runs keep af.
        xml_af = fn if contains_cjk(text) else af
        sz = size_pt if size_pt is not None else 12.0
        cr = color_rgb or fonts_conf.get("color_rgb", "000000")

        run = paragraph.add_run(text)
        run.font.name = fn
        run.font.size = Pt(sz)
        run.font.bold = bold

        rPr = run._r.get_or_add_rPr()
        # python-docx may already have inserted w:rFonts when run.font.name was
        # assigned. Replace it so East Asian and Western mappings cannot conflict.
        for existing in rPr.xpath('w:rFonts'):
            rPr.remove(existing)
        rFonts = OxmlElement("w:rFonts")
        for slot, value in (("ascii", xml_af), ("hAnsi", xml_af), ("eastAsia", fn),
                            ("hint", font_hint or "eastAsia")):
            rFonts.set(qn("w:" + slot), value)
        rPr.insert(0, rFonts)
        if cr:
            run.font.color.rgb = RGBColor.from_string(cr)
        return run

    def is_subhead(self, line_text, custom_keywords=None):
        """Check if a line should be automatically bolded as a subheading."""
        stripped = line_text.strip()
        keywords = custom_keywords or self.get_style("lesson_style", "subhead_bold_keywords", default=[])
        if stripped.startswith("【"):
            return True
        for kw in keywords:
            if stripped == kw or stripped.startswith(kw):
                return True
        return False

    def add_cell_text(self, paragraph, text, font_name=None, size_pt=None, bold=False):
        """Style recognized technical tokens without touching source characters.

        Headings/labels retain their configured font. Only ASCII font slots of
        technical runs change; CJK text, size and weight retain the body style.
        This improves readability, not an absolute guarantee against wrapping
        at punctuation or when a token is wider than the available column.
        """
        conf = self.get_style("technical_text", default={})
        if bold or not conf.get("enabled", True):
            return [self.add_run(paragraph, text, font_name=font_name, size_pt=size_pt, bold=bold)]
        runs = []
        has_technical = False
        for chunk, technical in split_technical_text(text):
            has_technical |= technical
            runs.append(self.add_run(
                paragraph, chunk, font_name=font_name, size_pt=size_pt, bold=bold,
                ascii_font=conf.get("font_name", "Courier New") if technical else None,
                font_hint="default" if technical else None,
            ))
            if technical:
                r_pr = runs[-1]._r.get_or_add_rPr()
                lang = OxmlElement("w:lang")
                lang.set(qn("w:val"), "en-US")
                lang.set(qn("w:eastAsia"), "zh-CN")
                r_pr.append(lang)
        if has_technical:
            ppr = paragraph._p.get_or_add_pPr()
            # CT_PPr schema order: these fields precede spacing/ind/jc. Word's
            # ordinary word-level wrapping still permits punctuation breaks.
            successors = (
                "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
                "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
                "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
                "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
            )
            for tag, following in (
                ("wordWrap", ("w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN") + successors),
                ("autoSpaceDE", ("w:autoSpaceDN",) + successors),
                ("autoSpaceDN", successors),
                ("snapToGrid", ("w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
                                 "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
                                 "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle",
                                 "w:rPr", "w:sectPr", "w:pPrChange")),
            ):
                for old in ppr.findall(qn("w:" + tag)):
                    ppr.remove(old)
                setting = OxmlElement("w:" + tag)
                # wordWrap is counter-intuitive in WordprocessingML: on means
                # wrap at word boundaries; off permits character-level splits.
                setting.set(qn("w:val"), "1" if tag == "wordWrap" else "0")
                ppr.insert_element_before(setting, *following)
        return runs

    def format_cell(self, cell, content, align=None, bold=False, size_pt=None, font_name=None,
                    line_spacing=None, space_before=None, space_after=None,
                    subhead_keywords=None, style_override=None):
        """
        Format a lesson table cell with vertical center alignment, paragraph spacing,
        subhead detection, and optional override.
        """
        cell_conf = copy.deepcopy(self.get_style("lesson_style", "cell_format", default={}))
        if style_override:
            cell_conf = deep_merge(cell_conf, style_override)

        fn = font_name or cell_conf.get("font_name", "宋体")
        sz = size_pt if size_pt is not None else cell_conf.get("size_pt", 12.0)
        lsp = line_spacing if line_spacing is not None else cell_conf.get("line_spacing", 1.15)
        sb = space_before if space_before is not None else cell_conf.get("space_before_pt", 1.5)
        sa = space_after if space_after is not None else cell_conf.get("space_after_pt", 1.5)
        al = self.parse_alignment(align) if align is not None else WD_ALIGN_PARAGRAPH.LEFT

        cell.text = ""
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
        tcPr.append(vAlign)

        if isinstance(content, str):
            lines = content.split('\n')
            for i, line in enumerate(lines):
                p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                p.alignment = al
                p.paragraph_format.line_spacing = lsp
                p.paragraph_format.space_before = Pt(sb)
                p.paragraph_format.space_after = Pt(sa)
                is_sub = bold or self.is_subhead(line, subhead_keywords)
                self.add_cell_text(p, line, font_name=fn, size_pt=sz, bold=is_sub)
        elif isinstance(content, list):
            for i, item in enumerate(content):
                p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
                p.paragraph_format.line_spacing = lsp
                p.paragraph_format.space_before = Pt(sb)
                p.paragraph_format.space_after = Pt(sa)
                p.alignment = al

                if isinstance(item, tuple) or isinstance(item, list):
                    t = item[0]
                    b = item[1] if len(item) > 1 else bold
                    item_fn = item[2] if len(item) > 2 else fn
                    item_sz = item[3] if len(item) > 3 else sz
                    self.add_cell_text(p, t, font_name=item_fn, size_pt=item_sz, bold=b)
                elif isinstance(item, dict):
                    t = item.get("text", "")
                    b = item.get("bold", bold)
                    item_fn = item.get("font_name", fn)
                    item_sz = item.get("size_pt", sz)
                    self.add_cell_text(p, t, font_name=item_fn, size_pt=item_sz, bold=b)
                else:
                    self.add_cell_text(p, str(item), font_name=fn, size_pt=sz, bold=bold)

        # Keep individual paragraphs readable across pages, not entire process
        # rows (which may be longer than one page). A heading chain ends at the
        # first body paragraph, and never links into the next cell or row.
        paragraphs = cell.paragraphs
        for index, paragraph in enumerate(paragraphs):
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.widow_control = True
            text_runs = [run for run in paragraph.runs if run.text.strip()]
            is_heading = self.is_subhead(paragraph.text, subhead_keywords) or (
                bool(text_runs) and all(run.bold for run in text_runs)
            )
            paragraph.paragraph_format.keep_with_next = bool(
                paragraph.text.strip() and is_heading and index < len(paragraphs) - 1
            )

    def format_prog_cell(self, cell, content, align=None, bold=False, size_pt=None, font_name=None,
                         line_spacing=None, space_before=None, space_after=None, style_override=None):
        """Format a schedule table cell."""
        prog_conf = copy.deepcopy(self.get_style("schedule_style", "cell_format", default={}))
        if style_override:
            prog_conf = deep_merge(prog_conf, style_override)

        fn = font_name or prog_conf.get("font_name", "宋体")
        sz = size_pt if size_pt is not None else prog_conf.get("size_pt", 12.0)
        lsp = line_spacing if line_spacing is not None else prog_conf.get("line_spacing", 1.1)
        sb = space_before if space_before is not None else prog_conf.get("space_before_pt", 0.0)
        sa = space_after if space_after is not None else prog_conf.get("space_after_pt", 0.0)
        al = self.parse_alignment(align) if align is not None else WD_ALIGN_PARAGRAPH.LEFT

        cell.text = ""
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
        tcPr.append(vAlign)

        lines = content.split('\n') if isinstance(content, str) else [str(content)]
        for i, line in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = al
            p.paragraph_format.line_spacing = lsp
            p.paragraph_format.space_before = Pt(sb)
            p.paragraph_format.space_after = Pt(sa)
            self.add_cell_text(p, line, font_name=fn, size_pt=sz, bold=bold)

    def setup_table_structure(self, table, table_config):
        """Clean and insert tblPr and tblGrid XML for custom table geometry."""
        for old_pr in table._element.xpath('w:tblPr'):
            table._element.remove(old_pr)
        for old_grid in table._element.xpath('w:tblGrid'):
            table._element.remove(old_grid)

        tbl_style = table_config.get("tbl_style", "5")
        cell_mar = table_config.get("cell_margins_dxa", {"top": 40, "left": 108, "bottom": 40, "right": 108})
        grid_cols = table_config.get("grid_cols_dxa", [])
        total_width = sum(int(width) for width in grid_cols)

        # Keep the table geometry deterministic across Word-compatible editors.
        # The printable width is calculated from the same page settings used by
        # apply_page_setup(); a narrow table is centered with an explicit indent.
        page_conf = self.get_style("page", default={})
        printable_width = round(
            (
                page_conf.get("page_width_pt", 595.3)
                - page_conf.get("left_margin_pt", 84.6)
                - page_conf.get("right_margin_pt", 84.6)
            ) * 20
        )
        table_indent = table_config.get(
            "table_indent_dxa",
            max(0, (printable_width - total_width) // 2) if total_width else 0,
        )

        tblPr_xml = parse_xml(f'''
        <w:tblPr {nsdecls("w")}>
          <w:tblStyle w:val="{tbl_style}"/>
          <w:tblW w:w="{total_width}" w:type="dxa"/>
          <w:tblInd w:w="{table_indent}" w:type="dxa"/>
          <w:tblBorders>
            <w:top w:val="single" w:color="000000" w:sz="4" w:space="0"/>
            <w:left w:val="single" w:color="000000" w:sz="4" w:space="0"/>
            <w:bottom w:val="single" w:color="000000" w:sz="4" w:space="0"/>
            <w:right w:val="single" w:color="000000" w:sz="4" w:space="0"/>
            <w:insideH w:val="single" w:color="000000" w:sz="4" w:space="0"/>
            <w:insideV w:val="single" w:color="000000" w:sz="4" w:space="0"/>
          </w:tblBorders>
          <w:tblLayout w:type="fixed"/>
          <w:tblCellMar>
            <w:top w:w="{cell_mar.get('top', 40)}" w:type="dxa"/>
            <w:left w:w="{cell_mar.get('left', 108)}" w:type="dxa"/>
            <w:bottom w:w="{cell_mar.get('bottom', 40)}" w:type="dxa"/>
            <w:right w:w="{cell_mar.get('right', 108)}" w:type="dxa"/>
          </w:tblCellMar>
        </w:tblPr>
        ''')
        table._element.insert(0, tblPr_xml)

        if grid_cols:
            cols_xml_str = "".join([f'<w:gridCol w:w="{col_w}"/>' for col_w in grid_cols])
            tblGrid = parse_xml(f'''
            <w:tblGrid {nsdecls("w")}>
              {cols_xml_str}
            </w:tblGrid>
            ''')
            table._element.insert(1, tblGrid)

    def apply_dynamic_cell_widths(self, table, grid_cols_dxa, row_spans_list=None, cant_split=False):
        """
        Dynamically calculate and apply exact XML cell widths (w:tcW) and cantSplit
        for every cell in the table based on grid_cols_dxa and row column spans.
        Eliminates the need for hardcoded row index maps.
        """
        for r_i, row in enumerate(table.rows):
            if cant_split:
                trPr = row._tr.get_or_add_trPr()
                if not trPr.xpath('w:cantSplit'):
                    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

            if row_spans_list and r_i < len(row_spans_list):
                spans = row_spans_list[r_i]
            else:
                spans = [(c, c) for c in range(len(grid_cols_dxa))]

            # row.cells expands horizontally merged cells into repeated proxy
            # objects. Work on physical <w:tc> elements instead so each tcW is
            # written once and equals exactly the grid columns represented by it.
            physical_cells = list(row._tr.tc_lst)
            if len(physical_cells) == len(spans):
                physical_spans = spans
            else:
                # Defensive fallback for malformed/custom templates. Derive the
                # physical spans from gridSpan, maintaining a running grid cursor.
                physical_spans = []
                grid_cursor = 0
                for tc in physical_cells:
                    tc_pr = tc.get_or_add_tcPr()
                    grid_span = tc_pr.xpath('w:gridSpan')
                    span_size = int(grid_span[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')) if grid_span else 1
                    physical_spans.append((grid_cursor, grid_cursor + span_size - 1))
                    grid_cursor += span_size

            for tc, (start_col, end_col) in zip(physical_cells, physical_spans):
                tcPr_cell = tc.get_or_add_tcPr()
                for old_w in tcPr_cell.xpath('w:tcW'):
                    tcPr_cell.remove(old_w)
                w_val = sum(grid_cols_dxa[start_col : end_col + 1])
                w_xml = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{w_val}" w:type="dxa"/>')
                tcPr_cell.insert(0, w_xml)

    def set_repeat_table_header(self, row):
        """Mark a table row as a repeating header in Word-compatible editors."""
        tr_pr = row._tr.get_or_add_trPr()
        if not tr_pr.xpath('w:tblHeader'):
            tr_pr.append(parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>'))

    def set_row_cant_split(self, row):
        """Keep a short structural row together without forcing all content rows."""
        tr_pr = row._tr.get_or_add_trPr()
        if not tr_pr.xpath('w:cantSplit'):
            cant_split = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
            # Preserve the conventional CT_TrPr property order.
            header_nodes = tr_pr.xpath('w:tblHeader')
            if header_nodes:
                tr_pr.insert(tr_pr.index(header_nodes[0]), cant_split)
            else:
                tr_pr.append(cant_split)

    def apply_cell_widths(self, table, widths_map, cant_split=False):
        """Apply explicit cell widths (dxa) and optional cantSplit properties to all table rows/cells."""
        for r_i, row in enumerate(table.rows):
            if cant_split:
                trPr = row._tr.get_or_add_trPr()
                trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

            row_widths = widths_map.get(str(r_i)) or widths_map.get(r_i) or widths_map.get("default_row")
            if not row_widths:
                continue

            for c_i, cell in enumerate(row.cells):
                tcPr_cell = cell._tc.get_or_add_tcPr()
                for old_w in tcPr_cell.xpath('w:tcW'):
                    tcPr_cell.remove(old_w)
                w_val = row_widths[c_i] if c_i < len(row_widths) else row_widths[-1]
                w_xml = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{w_val}" w:type="dxa"/>')
                tcPr_cell.insert(0, w_xml)

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
inspect_docx: 嗅探并逆向提取已有 Word (.docx) 教案的排版特征、样式参数与表格骨架。
"""

import sys
import os
import json
import argparse

# 1. 依赖优雅自检
try:
    import docx
    from docx.shared import Inches, Pt, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("❌ [环境提示] 未检测到 Word 文档处理库 'python-docx'。", file=sys.stderr)
    print("👉 请在当前 Python 环境中运行: pip install python-docx", file=sys.stderr)
    sys.exit(2)


def dxa_to_pt(dxa_val):
    if dxa_val is None:
        return None
    return round(dxa_val / 20.0, 2)


def dxa_to_cm(dxa_val):
    if dxa_val is None:
        return None
    return round(dxa_val / 567.0, 2)


def inspect_document(docx_path):
    """分析 Word 文档排版属性与结构骨架。"""
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"未找到文档文件: {docx_path}")

    doc = docx.Document(docx_path)
    result = {
        "file_info": {
            "path": os.path.abspath(docx_path),
            "filename": os.path.basename(docx_path),
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "total_sections": len(doc.sections)
        },
        "page_layout": {},
        "font_hierarchy": {},
        "tables_summary": [],
        "document_structure_skeleton": []
    }

    # 1. 页面布局与边距
    if doc.sections:
        sec = doc.sections[0]
        result["page_layout"] = {
            "page_width_cm": dxa_to_cm(sec.page_width.twips if hasattr(sec.page_width, "twips") else None),
            "page_height_cm": dxa_to_cm(sec.page_height.twips if hasattr(sec.page_height, "twips") else None),
            "orientation": "LANDSCAPE" if sec.orientation == docx.enum.section.WD_ORIENT.LANDSCAPE else "PORTRAIT",
            "margins_cm": {
                "top": dxa_to_cm(sec.top_margin.twips if hasattr(sec.top_margin, "twips") else None),
                "bottom": dxa_to_cm(sec.bottom_margin.twips if hasattr(sec.bottom_margin, "twips") else None),
                "left": dxa_to_cm(sec.left_margin.twips if hasattr(sec.left_margin, "twips") else None),
                "right": dxa_to_cm(sec.right_margin.twips if hasattr(sec.right_margin, "twips") else None)
            },
            "margins_dxa": {
                "top": sec.top_margin.twips if hasattr(sec.top_margin, "twips") else 1440,
                "bottom": sec.bottom_margin.twips if hasattr(sec.bottom_margin, "twips") else 1440,
                "left": sec.left_margin.twips if hasattr(sec.left_margin, "twips") else 1800,
                "right": sec.right_margin.twips if hasattr(sec.right_margin, "twips") else 1800
            }
        }

    # 2. 统计所有段落与字号/字体分布
    fonts_detected = set()
    sizes_detected = set()
    sample_headings = []

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        p_fonts = set()
        p_sizes = set()

        for r in p.runs:
            # 尝试提取中文字体与西文字体
            if r.font.name:
                p_fonts.add(r.font.name)
            # 检查 w:rFonts
            rPr = r._r.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                eastAsia = rFonts.get(qn('w:eastAsia'))
                if eastAsia:
                    p_fonts.add(eastAsia)
            if r.font.size:
                p_sizes.add(round(r.font.size.pt, 1))

        fonts_detected.update(p_fonts)
        sizes_detected.update(p_sizes)

        # 记录前几个具有标题特征的段落
        if len(text) < 40 and (p_sizes or len(sample_headings) < 8):
            sample_headings.append({
                "paragraph_index": i,
                "text": text,
                "fonts": list(p_fonts),
                "sizes_pt": list(p_sizes),
                "alignment": str(p.alignment) if p.alignment else "DEFAULT"
            })

    result["font_hierarchy"] = {
        "all_detected_fonts": sorted(list(fonts_detected)),
        "all_detected_sizes_pt": sorted(list(sizes_detected)),
        "sample_headings": sample_headings[:10]
    }

    # 3. 分析表格结构
    for t_idx, table in enumerate(doc.tables):
        rows_cnt = len(table.rows)
        cols_cnt = len(table.columns) if rows_cnt > 0 else 0

        # 分析各列大致宽度 (从第一行单元格提取)
        col_widths_dxa = []
        if rows_cnt > 0:
            for cell in table.rows[0].cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None and tcW.get(qn('w:w')):
                    try:
                        col_widths_dxa.append(int(tcW.get(qn('w:w'))))
                    except ValueError:
                        pass

        # 提取前几行的表头样本与内容特征
        table_cells_sample = []
        for r_i in range(min(rows_cnt, 8)):
            row_cells = table.rows[r_i].cells
            row_texts = []
            seen_texts = set()
            for c in row_cells:
                t = c.text.strip().replace("\n", " ")
                # 简单过滤合并单元格重复内容
                if t not in seen_texts:
                    seen_texts.add(t)
                    row_texts.append(t[:30] + ("..." if len(t) > 30 else ""))
            table_cells_sample.append({
                "row_index": r_i,
                "cell_contents": row_texts
            })

        # 判断表格类型（进度表、封面信息表、课时教学设计表等）
        first_row_text = " ".join([c.text for c in table.rows[0].cells]) if rows_cnt > 0 else ""
        table_type = "generic_table"
        if "授课计划" in first_row_text or "进度" in first_row_text or ("周次" in first_row_text and "课时" in first_row_text):
            table_type = "schedule_table"
        elif "课程名称" in first_row_text or "任课教师" in first_row_text or "班级" in first_row_text:
            table_type = "cover_or_meta_table"
        elif any("环节" in c.text or "教学过程" in c.text or "教学活动" in c.text for c in (table.rows[min(3, rows_cnt-1)].cells if rows_cnt > 3 else [])):
            table_type = "lesson_design_table"

        result["tables_summary"].append({
            "table_index": t_idx,
            "inferred_type": table_type,
            "rows_count": rows_cnt,
            "columns_count": cols_cnt,
            "col_widths_dxa": col_widths_dxa,
            "sample_rows": table_cells_sample
        })

    # 4. 文档骨架概览（段落与表格出现的流式顺序）
    body_elements = []
    p_idx = 0
    t_idx = 0
    for child in doc.element.body:
        if child.tag.endswith('p'):
            if p_idx < len(doc.paragraphs):
                p_text = doc.paragraphs[p_idx].text.strip()
                if p_text:
                    body_elements.append({
                        "type": "paragraph",
                        "index": p_idx,
                        "snippet": p_text[:40] + ("..." if len(p_text) > 40 else "")
                    })
                p_idx += 1
        elif child.tag.endswith('tbl'):
            if t_idx < len(doc.tables):
                tbl_info = result["tables_summary"][t_idx]
                body_elements.append({
                    "type": "table",
                    "index": t_idx,
                    "inferred_type": tbl_info["inferred_type"],
                    "rows": tbl_info["rows_count"],
                    "cols": tbl_info["columns_count"]
                })
                t_idx += 1

    result["document_structure_skeleton"] = body_elements[:25]
    return result


def main():
    parser = argparse.ArgumentParser(
        description="Word (.docx) 教案排版嗅探与骨架提取工具",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument("docx_path", help="待分析的 Word (.docx) 教案文件路径")
    parser.add_argument("-o", "--output", help="将分析结果保存为指定 JSON 文件（默认输出到控制台）")
    parser.add_argument("-q", "--quiet", action="store_true", help="静默模式，仅输出最终 JSON")

    args = parser.parse_args()

    try:
        data = inspect_document(args.docx_path)
        json_output = json.dumps(data, ensure_ascii=False, indent=2)

        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(json_output)
            if not args.quiet:
                print(f"✓ 分析结果已保存至: {args.output}")
        else:
            print(json_output)

    except Exception as e:
        print(f"❌ [错误] 分析失败: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()

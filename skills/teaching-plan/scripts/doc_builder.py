# -*- coding: utf-8 -*-
"""
TeachingPlanDocBuilder: Builds Word (.docx) documents from structured content data
using declarative layout templates and typography configurations.
"""

import sys
try:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
except ImportError:
    print("❌ [环境提示] 未检测到 Word 文档处理库 'python-docx'。", file=sys.stderr)
    print("👉 请在当前 Python 环境中运行: pip install python-docx", file=sys.stderr)
    sys.exit(2)

import os
import copy
import math
import warnings
# Ensure sibling modules can be imported when called from any working directory
sys.path.insert(0, os.path.dirname(__file__))

from style_manager import StyleManager
from template_manager import TemplateManager
from data_normalizer import DataNormalizer
from pagination import split_paragraph_blocks, balance_cell_blocks, estimate_text_height

class TeachingPlanDocBuilder:
    def __init__(self, content_data, template=None, style_manager=None):
        """
        :param content_data: Dictionary containing 'cover', 'schedule', and 'lessons' data.
        :param template: Instance of TemplateManager, path to template JSON, or template dict.
        :param style_manager: Instance of StyleManager. If None, default StyleManager is instantiated.
        """
        self.content = DataNormalizer.normalize_content(content_data or {})
        if isinstance(template, TemplateManager):
            self.tm = template
        else:
            self.tm = TemplateManager(template)
        self.sm = style_manager or StyleManager()

    def build_document(self):
        """Construct the complete docx.Document."""
        doc = docx.Document()

        # 1. Page Setup
        section = doc.sections[0]
        self.sm.apply_page_setup(section)

        # 2. Cover Page
        cover_data = self.content.get("cover")
        has_prior_section = False
        if cover_data:
            self._build_cover_page(doc, cover_data)
            has_prior_section = True

        # 3. Schedule Table
        schedule_data = self.content.get("schedule")
        if schedule_data:
            self._build_schedule_page(doc, schedule_data, page_break_before=has_prior_section)
            has_prior_section = True

        # 4. Lesson Plans
        lessons_data = self.content.get("lessons", [])
        total_lessons = len(lessons_data)
        pagination = self.tm.get_section("pagination", default={}) or {}
        lesson_start = pagination.get("lesson_start", "new_page")
        if lesson_start not in {"new_page", "flow"}:
            raise ValueError("pagination.lesson_start 必须是 new_page 或 flow")
        for idx, lesson in enumerate(lessons_data):
            self._build_lesson_table(
                doc,
                lesson,
                idx,
                total_lessons,
                page_break_before=has_prior_section and (idx == 0 or lesson_start == "new_page"),
            )
            has_prior_section = True

        return doc

    def _build_cover_page(self, doc, cover):
        cover_conf = self.sm.get_style("cover_style", default={})
        tmpl_cover = self.tm.get_section("cover", default={})
        context = {"cover": cover}

        # Top spacing
        top_sp = tmpl_cover.get("top_spacing_paragraphs", cover_conf.get("top_spacing_paragraphs", 5))
        for _ in range(top_sp):
            doc.add_paragraph()

        # Title
        t_conf = cover_conf.get("title", {})
        tmpl_t = tmpl_cover.get("title", {})
        title_text = self.tm.evaluate_expression(tmpl_t.get("bind", "{{ cover.title | default('教  学  设  计') }}"), context)
        p_title = doc.add_paragraph()
        p_title.alignment = self.sm.parse_alignment(tmpl_t.get("align", t_conf.get("align", "CENTER")))
        self.sm.add_run(
            p_title,
            title_text,
            font_name=tmpl_t.get("font_name", t_conf.get("font_name", "微软雅黑")),
            size_pt=tmpl_t.get("size_pt", t_conf.get("size_pt", 42.0)),
            bold=tmpl_t.get("bold", t_conf.get("bold", True))
        )

        # Subtitle
        sub_conf = cover_conf.get("subtitle", {})
        tmpl_sub = tmpl_cover.get("subtitle", {})
        subtitle_text = self.tm.evaluate_expression(tmpl_sub.get("bind", "{{ cover.subtitle | default('') }}"), context)
        p_sub = doc.add_paragraph()
        p_sub.alignment = self.sm.parse_alignment(tmpl_sub.get("align", sub_conf.get("align", "CENTER")))
        self.sm.add_run(
            p_sub,
            subtitle_text,
            font_name=tmpl_sub.get("font_name", sub_conf.get("font_name", "楷体_GB2312")),
            size_pt=tmpl_sub.get("size_pt", sub_conf.get("size_pt", 16.0)),
            bold=tmpl_sub.get("bold", sub_conf.get("bold", False))
        )

        # Mid spacing
        mid_sp = tmpl_cover.get("mid_spacing_paragraphs", cover_conf.get("mid_spacing_paragraphs", 4))
        for _ in range(mid_sp):
            doc.add_paragraph()

        # Info Table
        info_items = cover.get("info_items", [])
        if info_items:
            tbl_info_conf = tmpl_cover.get("info_table") or cover_conf.get("info_table", {})
            col_widths = tbl_info_conf.get("col_widths_dxa", [2000, 3800])
            total_w = sum(col_widths)
            cell_mar = tbl_info_conf.get("cell_margins_dxa", {"top": 40, "bottom": 40, "left": 0, "right": 0})
            lbl_style = tbl_info_conf.get("label_style", {})
            val_style = tbl_info_conf.get("val_style", {})

            t_info = doc.add_table(rows=len(info_items), cols=2)

            for old_pr in t_info._element.xpath('w:tblPr'):
                t_info._element.remove(old_pr)
            for old_grid in t_info._element.xpath('w:tblGrid'):
                t_info._element.remove(old_grid)

            tblPr_info = parse_xml(f'''
            <w:tblPr {nsdecls("w")}>
              <w:jc w:val="center"/>
              <w:tblW w:w="{total_w}" w:type="dxa"/>
              <w:tblBorders>
                <w:top w:val="none"/>
                <w:left w:val="none"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
                <w:insideH w:val="none"/>
                <w:insideV w:val="none"/>
              </w:tblBorders>
              <w:tblLayout w:type="fixed"/>
              <w:tblCellMar>
                <w:top w:w="{cell_mar.get('top', 40)}" w:type="dxa"/>
                <w:left w:w="{cell_mar.get('left', 0)}" w:type="dxa"/>
                <w:bottom w:w="{cell_mar.get('bottom', 40)}" w:type="dxa"/>
                <w:right w:w="{cell_mar.get('right', 0)}" w:type="dxa"/>
              </w:tblCellMar>
            </w:tblPr>
            ''')
            t_info._element.insert(0, tblPr_info)

            tblGrid_info = parse_xml(f'''
            <w:tblGrid {nsdecls("w")}>
              <w:gridCol w:w="{col_widths[0]}"/>
              <w:gridCol w:w="{col_widths[1]}"/>
            </w:tblGrid>
            ''')
            t_info._element.insert(1, tblGrid_info)

            for row_idx, item in enumerate(info_items):
                label, val = item[0], item[1]
                cell_lbl = t_info.cell(row_idx, 0)
                cell_val = t_info.cell(row_idx, 1)

                tcPr_0 = cell_lbl._tc.get_or_add_tcPr()
                for old_w in tcPr_0.xpath('w:tcW'):
                    tcPr_0.remove(old_w)
                tcPr_0.append(parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[0]}" w:type="dxa"/>'))

                tcPr_1 = cell_val._tc.get_or_add_tcPr()
                for old_w in tcPr_1.xpath('w:tcW'):
                    tcPr_1.remove(old_w)
                tcPr_1.append(parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[1]}" w:type="dxa"/>'))

                p_lbl = cell_lbl.paragraphs[0]
                p_lbl.alignment = self.sm.parse_alignment(lbl_style.get("align", "LEFT"))
                p_lbl.paragraph_format.line_spacing = lbl_style.get("line_spacing", 1.3)
                p_lbl.paragraph_format.space_before = Pt(lbl_style.get("space_before_pt", 4.0))
                p_lbl.paragraph_format.space_after = Pt(lbl_style.get("space_after_pt", 4.0))
                self.sm.add_run(
                    p_lbl, label,
                    font_name=lbl_style.get("font_name", "楷体_GB2312"),
                    size_pt=lbl_style.get("size_pt", 16.0),
                    bold=lbl_style.get("bold", True)
                )

                p_val = cell_val.paragraphs[0]
                p_val.alignment = self.sm.parse_alignment(val_style.get("align", "LEFT"))
                p_val.paragraph_format.line_spacing = val_style.get("line_spacing", 1.3)
                p_val.paragraph_format.space_before = Pt(val_style.get("space_before_pt", 4.0))
                p_val.paragraph_format.space_after = Pt(val_style.get("space_after_pt", 4.0))
                self.sm.add_run(
                    p_val, val,
                    font_name=val_style.get("font_name", "楷体_GB2312"),
                    size_pt=val_style.get("size_pt", 16.0),
                    bold=val_style.get("bold", False)
                )

        # Bottom spacing
        bot_sp = tmpl_cover.get("bottom_spacing_paragraphs", cover_conf.get("bottom_spacing_paragraphs", 5))
        for _ in range(bot_sp):
            doc.add_paragraph()

        # School
        sch_conf = cover_conf.get("school", {})
        tmpl_sch = tmpl_cover.get("school", {})
        sch_text = self.tm.evaluate_expression(tmpl_sch.get("bind", "{{ cover.school | default('') }}"), context)
        p_school = doc.add_paragraph()
        p_school.alignment = self.sm.parse_alignment(tmpl_sch.get("align", sch_conf.get("align", "CENTER")))
        p_school.paragraph_format.space_after = Pt(tmpl_sch.get("space_after_pt", sch_conf.get("space_after_pt", 6.0)))
        self.sm.add_run(
            p_school,
            sch_text,
            font_name=tmpl_sch.get("font_name", sch_conf.get("font_name", "微软雅黑")),
            size_pt=tmpl_sch.get("size_pt", sch_conf.get("size_pt", 16.0)),
            bold=tmpl_sch.get("bold", sch_conf.get("bold", True))
        )

        # Date
        date_conf = cover_conf.get("date", {})
        tmpl_date = tmpl_cover.get("date", {})
        date_text = self.tm.evaluate_expression(tmpl_date.get("bind", "{{ cover.date | default('') }}"), context)
        p_date = doc.add_paragraph()
        p_date.alignment = self.sm.parse_alignment(tmpl_date.get("align", date_conf.get("align", "CENTER")))
        p_date.paragraph_format.space_after = Pt(tmpl_date.get("space_after_pt", date_conf.get("space_after_pt", 0.0)))
        self.sm.add_run(
            p_date,
            date_text,
            font_name=tmpl_date.get("font_name", date_conf.get("font_name", "微软雅黑")),
            size_pt=tmpl_date.get("size_pt", date_conf.get("size_pt", 14.0)),
            bold=tmpl_date.get("bold", date_conf.get("bold", False))
        )

    @staticmethod
    def _split_table_at_rows(table, split_rows, continuous_process_borders=False, keep_before_rows=()):
        """Split a table before the requested rows while preserving its geometry.

        Teaching-process headings can only repeat reliably when they are at the
        start of their table. Word silently coalesces immediately adjacent tables,
        which turns an internal ``tblHeader`` row back into a non-header row. A
        one-point separator paragraph therefore keeps the metadata, process, and
        reflection fragments semantically independent in Word and WPS.
        """
        source = table._tbl
        all_rows = list(source.tr_lst)
        # Only explicit short-tail anchors may bridge a table boundary. Linking
        # every process row would pull large blocks onto the following page.
        for boundary in keep_before_rows:
            if 0 < boundary < len(all_rows):
                for p in all_rows[boundary - 1].xpath("w:tc/w:p"):
                    p.get_or_add_pPr().get_or_add_keepNext().val = True
        boundaries = [0]
        boundaries.extend(sorted({int(row) for row in split_rows if 0 < int(row) < len(all_rows)}))
        boundaries.append(len(all_rows))
        if len(boundaries) <= 2:
            return

        parent = source.getparent()
        insert_at = parent.index(source)
        fragments = []
        for start, end in zip(boundaries, boundaries[1:]):
            fragment = copy.deepcopy(source)
            for row_idx, row_element in reversed(list(enumerate(fragment.tr_lst))):
                if row_idx < start or row_idx >= end:
                    fragment.remove(row_element)
            if continuous_process_borders and fragment.tr_lst[0].xpath("w:trPr/w:tblHeader"):
                # Let the table's own outer border close each printed page.
                # Per-cell nil borders would suppress that continuation border.
                for border in fragment.xpath("w:tblPr/w:tblBorders/w:insideH"):
                    border.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "nil")
            fragments.append(fragment)

        parent.remove(source)
        cursor = insert_at
        for fragment_idx, fragment in enumerate(fragments):
            parent.insert(cursor, fragment)
            cursor += 1
            if fragment_idx < len(fragments) - 1:
                separator = parse_xml(f'''
                    <w:p {nsdecls("w")}>
                      <w:pPr>
                        <w:spacing w:before="0" w:after="0" w:line="20" w:lineRule="exact"/>
                      </w:pPr>
                    </w:p>
                ''')
                if boundaries[fragment_idx + 1] in keep_before_rows:
                    separator.get_or_add_pPr().get_or_add_keepNext().val = True
                parent.insert(cursor, separator)
                cursor += 1

    def _build_schedule_page(self, doc, schedule, page_break_before=False):
        sch_style = self.sm.get_style("schedule_style", default={})
        tmpl_sch = self.tm.get_section("schedule", default={})
        context = {"schedule": schedule}

        # Schedule Title
        t_conf = sch_style.get("title", {})
        tmpl_t = tmpl_sch.get("title", {})
        title_text = self.tm.evaluate_expression(tmpl_t.get("bind", "{{ schedule.title | default('教学进度表') }}"), context)
        p_title = doc.add_paragraph()
        if page_break_before:
            p_title.paragraph_format.page_break_before = True
        p_title.paragraph_format.keep_with_next = True
        p_title.alignment = self.sm.parse_alignment(tmpl_t.get("align", t_conf.get("align", "CENTER")))
        p_title.paragraph_format.space_before = Pt(tmpl_t.get("space_before_pt", t_conf.get("space_before_pt", 0.0)))
        p_title.paragraph_format.space_after = Pt(tmpl_t.get("space_after_pt", t_conf.get("space_after_pt", 6.0)))
        self.sm.add_run(
            p_title,
            title_text,
            font_name=tmpl_t.get("font_name", t_conf.get("font_name", "微软雅黑")),
            size_pt=tmpl_t.get("size_pt", t_conf.get("size_pt", 18.0)),
            bold=tmpl_t.get("bold", t_conf.get("bold", True))
        )

        rows_data = schedule.get("rows", [])
        stats = schedule.get("stats", {})
        notes = schedule.get("notes", "")

        # Table configuration
        tbl_conf = tmpl_sch.get("table") or sch_style.get("table", {})
        grid_cols_dxa = tbl_conf.get("grid_cols_dxa", [650, 500, 400, 850, 1750, 650, 750, 1350, 972, 650])
        total_rows_count = 3 + len(rows_data) + 1
        t_prog = doc.add_table(rows=total_rows_count, cols=len(grid_cols_dxa))
        self.sm.setup_table_structure(t_prog, tbl_conf)

        row_spans = []

        # Row 0
        t_prog.cell(0, 0).merge(t_prog.cell(0, 2))
        self.sm.format_prog_cell(t_prog.cell(0, 0), "上课周数", align="CENTER", bold=True, size_pt=12.0)
        self.sm.format_prog_cell(t_prog.cell(0, 3), stats.get("weeks_count", "20"), align="CENTER", size_pt=12.0)
        self.sm.format_prog_cell(t_prog.cell(0, 4), "周学时", align="CENTER", bold=True, size_pt=12.0)
        t_prog.cell(0, 5).merge(t_prog.cell(0, 6))
        self.sm.format_prog_cell(t_prog.cell(0, 5), stats.get("week_hours", "6/班"), align="CENTER", size_pt=12.0)
        self.sm.format_prog_cell(t_prog.cell(0, 7), "总学时", align="CENTER", bold=True, size_pt=12.0)
        t_prog.cell(0, 8).merge(t_prog.cell(0, 9))
        self.sm.format_prog_cell(t_prog.cell(0, 8), stats.get("total_hours", "120"), align="CENTER", size_pt=12.0)
        row_spans.append([(0, 2), (3, 3), (4, 4), (5, 6), (7, 7), (8, 9)])

        # Row 1
        t_prog.cell(1, 0).merge(t_prog.cell(1, 2))
        self.sm.format_prog_cell(t_prog.cell(1, 0), "讲课学时", align="CENTER", bold=True, size_pt=12.0)
        self.sm.format_prog_cell(t_prog.cell(1, 3), stats.get("lecture_hours", "36"), align="CENTER", size_pt=12.0)
        self.sm.format_prog_cell(t_prog.cell(1, 4), "实验实作学时", align="CENTER", bold=True, size_pt=12.0)
        t_prog.cell(1, 5).merge(t_prog.cell(1, 6))
        self.sm.format_prog_cell(t_prog.cell(1, 5), stats.get("lab_hours", "72"), align="CENTER", size_pt=12.0)
        self.sm.format_prog_cell(t_prog.cell(1, 7), "复习考试", align="CENTER", bold=True, size_pt=12.0)
        t_prog.cell(1, 8).merge(t_prog.cell(1, 9))
        self.sm.format_prog_cell(t_prog.cell(1, 8), stats.get("exam_hours", "12"), align="CENTER", size_pt=12.0)
        row_spans.append([(0, 2), (3, 3), (4, 4), (5, 6), (7, 7), (8, 9)])

        # Row 2
        self.sm.format_prog_cell(t_prog.cell(2, 0), "周次", align="CENTER", bold=True, size_pt=12.0)
        t_prog.cell(2, 1).merge(t_prog.cell(2, 4))
        self.sm.format_prog_cell(t_prog.cell(2, 1), "理论内容", align="CENTER", bold=True, size_pt=12.0)
        self.sm.format_prog_cell(t_prog.cell(2, 5), "学时", align="CENTER", bold=True, size_pt=12.0)
        t_prog.cell(2, 6).merge(t_prog.cell(2, 8))
        self.sm.format_prog_cell(t_prog.cell(2, 6), "实作/实验内容", align="CENTER", bold=True, size_pt=12.0)
        self.sm.format_prog_cell(t_prog.cell(2, 9), "学时", align="CENTER", bold=True, size_pt=12.0)
        row_spans.append([(0, 0), (1, 4), (5, 5), (6, 8), (9, 9)])

        # Dynamic Schedule Rows
        for r_idx, row_item in enumerate(rows_data, start=3):
            w, th_txt, th_hrs, pr_txt, pr_hrs = row_item[0], row_item[1], row_item[2], row_item[3], row_item[4]
            self.sm.format_prog_cell(t_prog.cell(r_idx, 0), w, align="CENTER", size_pt=12.0)
            t_prog.cell(r_idx, 1).merge(t_prog.cell(r_idx, 4))
            self.sm.format_prog_cell(t_prog.cell(r_idx, 1), th_txt, align="LEFT", size_pt=12.0)
            self.sm.format_prog_cell(t_prog.cell(r_idx, 5), th_hrs, align="CENTER", size_pt=12.0)
            t_prog.cell(r_idx, 6).merge(t_prog.cell(r_idx, 8))
            self.sm.format_prog_cell(t_prog.cell(r_idx, 6), pr_txt, align="LEFT", size_pt=12.0)
            self.sm.format_prog_cell(t_prog.cell(r_idx, 9), pr_hrs, align="CENTER", size_pt=12.0)
            row_spans.append([(0, 0), (1, 4), (5, 5), (6, 8), (9, 9)])

        # Last row: 备注
        last_r = 3 + len(rows_data)
        t_prog.cell(last_r, 0).merge(t_prog.cell(last_r, 1))
        self.sm.format_prog_cell(t_prog.cell(last_r, 0), "备注", align="CENTER", bold=True, size_pt=12.0)
        t_prog.cell(last_r, 2).merge(t_prog.cell(last_r, 9))
        self.sm.format_prog_cell(t_prog.cell(last_r, 2), notes, align="LEFT", size_pt=12.0)
        row_spans.append([(0, 1), (2, 9)])

        # Apply exact calculated widths
        cant_split = tbl_conf.get("cant_split", True)
        self.sm.apply_dynamic_cell_widths(t_prog, grid_cols_dxa, row_spans_list=row_spans, cant_split=cant_split)
        # Word repeats only a contiguous header block at the start of a table;
        # include the two summary rows so the actual column headings remain valid.
        for header_row in t_prog.rows[:3]:
            self.sm.set_repeat_table_header(header_row)

    def _build_lesson_table(self, doc, lesson, idx, total_count, page_break_before=False):
        tmpl_lesson = self.tm.get_section("lesson_table", default={})
        lesson_style = self.sm.get_style("lesson_style", default={})
        lesson_override = lesson.get("style_override", {})
        context = {"lesson": lesson, "index": idx, "total": total_count}

        # Header Title
        tmpl_head = tmpl_lesson.get("title_header", {})
        head_conf = lesson_style.get("title_header", {})
        header_text = self.tm.evaluate_expression(tmpl_head.get("bind", "{{ lesson.lesson_num }}"), context)

        p_head = doc.add_paragraph()
        if page_break_before:
            p_head.paragraph_format.page_break_before = True
        p_head.paragraph_format.keep_with_next = True
        p_head.paragraph_format.space_before = Pt(tmpl_head.get("space_before_pt", head_conf.get("space_before_pt", 6.0)))
        p_head.paragraph_format.space_after = Pt(tmpl_head.get("space_after_pt", head_conf.get("space_after_pt", 2.0)))
        p_head.alignment = self.sm.parse_alignment(tmpl_head.get("align", head_conf.get("align", "LEFT")))
        self.sm.add_run(
            p_head,
            header_text,
            font_name=tmpl_head.get("font_name", head_conf.get("font_name", "微软雅黑")),
            size_pt=tmpl_head.get("size_pt", head_conf.get("size_pt", 14.0)),
            bold=tmpl_head.get("bold", head_conf.get("bold", True))
        )

        tbl_conf = tmpl_lesson.get("table") or lesson_style.get("table", {})
        grid_cols_dxa = tbl_conf.get("grid_cols_dxa", [1438, 1153, 4235, 1696])
        sections = tmpl_lesson.get("sections", [])

        # Expand only templates that opt into paragraph-sized process rows.
        # This makes Word's row-level pagination usable without a huge merged
        # stage cell or a change to the content JSON.
        expanded_repeaters = {}
        for section_index, sec in enumerate(sections):
            if sec.get("type") != "repeater":
                continue
            entries = []
            for stage_index, stage in enumerate(lesson.get("stages", [])):
                stage_context = {"lesson": lesson, "item": stage, "stage": stage, "stage_idx": stage_index}
                cell_blocks = []
                for cell_def in sec.get("row_template", {}).get("cells", []):
                    value = cell_def.get("text")
                    if value is None and "bind" in cell_def:
                        value = self.tm.evaluate_expression(cell_def["bind"], stage_context)
                    value = str(value or "")
                    blocks = (
                        split_paragraph_blocks(value, self.sm.is_subhead)
                        if sec.get("paragraph_groups") and not cell_def.get("bold")
                        else [value]
                    )
                    cell_blocks.append(blocks)
                count = max((len(blocks) for blocks in cell_blocks), default=1)
                if sec.get("paragraph_groups") and sec.get("balance_columns", False):
                    definitions = sec.get("row_template", {}).get("cells", [])
                    widths = []
                    offset = 0
                    for cell_def in definitions:
                        span = cell_def.get("colspan", 1)
                        widths.append(sum(grid_cols_dxa[offset:offset + span]) / 20 - 22)
                        offset += span
                    cell_format = self.sm.get_style("lesson_style", "cell_format", default={})
                    rows = balance_cell_blocks(cell_blocks, lambda column, text: estimate_text_height(
                        text, widths[column], definitions[column].get("size_pt", 12),
                        lesson_override.get("line_spacing", cell_format.get("line_spacing", 1.15)),
                        sec.get("paragraph_space_before_pt", lesson_override.get("space_before_pt", cell_format.get("space_before_pt", 1.5)))
                        + sec.get("paragraph_space_after_pt", lesson_override.get("space_after_pt", cell_format.get("space_after_pt", 1.5))),
                    ))
                    for block_index, cells in enumerate(rows):
                        entries.append((["\n".join(blocks) for blocks in cells], block_index, len(rows)))
                else:
                    for block_index in range(count):
                        entries.append(([
                            blocks[block_index] if block_index < len(blocks) else ""
                            for blocks in cell_blocks
                        ], block_index, count))
            expanded_repeaters[section_index] = entries

        # 1. Calculate required rows count dynamically
        total_rows = 0
        for section_index, sec in enumerate(sections):
            sec_type = sec.get("type")
            if sec_type == "static_row":
                total_rows += 1
            elif sec_type == "row_group":
                total_rows += len(sec.get("sub_rows", []))
            elif sec_type == "static_rows":
                total_rows += len(sec.get("items", []))
            elif sec_type == "repeater":
                total_rows += len(expanded_repeaters[section_index])
            elif sec_type == "reflection_blank":
                total_rows += 1

        t = doc.add_table(rows=total_rows, cols=len(grid_cols_dxa))
        self.sm.setup_table_structure(t, tbl_conf)

        current_row_idx = 0
        row_spans = []
        repeating_header_rows = []
        structural_rows = []
        oversized_rows = []
        split_rows = []
        reflection_starts = []
        group_edges = {}
        continuous_borders = any(sec.get("paragraph_groups") and sec.get("continuous_borders") for sec in sections)

        for section_index, sec in enumerate(sections):
            sec_type = sec.get("type")

            if sec_type == "static_row":
                r = current_row_idx
                current_row_idx += 1
                col_cursor = 0
                current_row_spans = []

                for cell_def in sec.get("cells", []):
                    c_span = cell_def.get("colspan", 1)
                    start_c = col_cursor
                    end_c = col_cursor + c_span - 1
                    col_cursor += c_span
                    current_row_spans.append((start_c, end_c))

                    if end_c > start_c:
                        t.cell(r, start_c).merge(t.cell(r, end_c))

                    cell_text = cell_def.get("text")
                    if cell_text is None and "bind" in cell_def:
                        cell_text = self.tm.evaluate_expression(cell_def["bind"], context)
                    cell_text = str(cell_text or "")

                    align = cell_def.get("align", "LEFT")
                    bold = cell_def.get("bold", False)
                    size_pt = cell_def.get("size_pt", 12.0)
                    self.sm.format_cell(
                        t.cell(r, start_c),
                        cell_text,
                        align=align,
                        bold=bold,
                        size_pt=size_pt,
                        style_override=lesson_override
                    )

                row_spans.append(current_row_spans)
                description = sec.get("description", "")
                if description in {"教学过程大标题", "教学过程表头"}:
                    repeating_header_rows.append(r)
                    structural_rows.append(r)
                    if description == "教学过程大标题":
                        split_rows.append(r)
                elif description == "教学反思大标题":
                    structural_rows.append(r)
                    split_rows.append(r)
                    reflection_starts.append(r)
                    for paragraph in t.rows[r].cells[0].paragraphs:
                        paragraph.paragraph_format.keep_with_next = True

            elif sec_type == "row_group":
                sub_rows = sec.get("sub_rows", [])
                group_len = len(sub_rows)
                start_r = current_row_idx
                end_r = start_r + group_len - 1

                # Group Header (Rowspan on Column 0)
                rowspan_hdr = sec.get("rowspan_header", {})
                hdr_text = rowspan_hdr.get("text", "")
                hdr_align = rowspan_hdr.get("align", "CENTER")
                hdr_bold = rowspan_hdr.get("bold", True)
                hdr_size = rowspan_hdr.get("size_pt", 12.0)

                for sub_i, sub_r_def in enumerate(sub_rows):
                    r = current_row_idx
                    current_row_idx += 1
                    col_cursor = 1  # Start at col 1 because col 0 is for the group header
                    current_row_spans = [(0, 0)]

                    for cell_def in sub_r_def.get("cells", []):
                        c_span = cell_def.get("colspan", 1)
                        start_c = col_cursor
                        end_c = col_cursor + c_span - 1
                        col_cursor += c_span
                        current_row_spans.append((start_c, end_c))

                        if end_c > start_c:
                            t.cell(r, start_c).merge(t.cell(r, end_c))

                        cell_text = cell_def.get("text")
                        if cell_text is None and "bind" in cell_def:
                            cell_text = self.tm.evaluate_expression(cell_def["bind"], context)
                        cell_text = str(cell_text or "")

                        align = cell_def.get("align", "LEFT")
                        bold = cell_def.get("bold", False)
                        size_pt = cell_def.get("size_pt", 12.0)
                        self.sm.format_cell(
                            t.cell(r, start_c),
                            cell_text,
                            align=align,
                            bold=bold,
                            size_pt=size_pt,
                            space_before=sec.get("paragraph_space_before_pt") if sec.get("paragraph_groups") else None,
                            space_after=sec.get("paragraph_space_after_pt") if sec.get("paragraph_groups") else None,
                            style_override=lesson_override
                        )

                    row_spans.append(current_row_spans)

                # Merge column 0 across group rows
                if group_len > 1:
                    t.cell(start_r, 0).merge(t.cell(end_r, 0))
                self.sm.format_cell(
                    t.cell(start_r, 0),
                    hdr_text,
                    align=hdr_align,
                    bold=hdr_bold,
                    size_pt=hdr_size,
                    style_override=lesson_override
                )

            elif sec_type == "static_rows":
                items = sec.get("items", [])
                row_template = sec.get("row_template", {})

                for item in items:
                    r = current_row_idx
                    current_row_idx += 1
                    col_cursor = 0
                    current_row_spans = []

                    # Pre-evaluate item dictionary values with current context
                    eval_item = {}
                    for ik, iv in item.items():
                        eval_item[ik] = self.tm.evaluate_expression(iv, context) if isinstance(iv, str) else iv
                    item_ctx = {"lesson": lesson, "item": eval_item, "index": idx}

                    for cell_def in row_template.get("cells", []):
                        c_span = cell_def.get("colspan", 1)
                        start_c = col_cursor
                        end_c = col_cursor + c_span - 1
                        col_cursor += c_span
                        current_row_spans.append((start_c, end_c))

                        if end_c > start_c:
                            t.cell(r, start_c).merge(t.cell(r, end_c))

                        cell_text = cell_def.get("text")
                        if cell_text is None and "bind" in cell_def:
                            cell_text = self.tm.evaluate_expression(cell_def["bind"], item_ctx)
                        cell_text = str(cell_text or "")

                        align = cell_def.get("align", "LEFT")
                        bold = cell_def.get("bold", False)
                        size_pt = cell_def.get("size_pt", 12.0)
                        self.sm.format_cell(
                            t.cell(r, start_c),
                            cell_text,
                            align=align,
                            bold=bold,
                            size_pt=size_pt,
                            style_override=lesson_override
                        )

                    row_spans.append(current_row_spans)

            elif sec_type == "repeater":
                row_template = sec.get("row_template", {})

                for cell_texts, block_index, block_count in expanded_repeaters[section_index]:
                    r = current_row_idx
                    current_row_idx += 1
                    col_cursor = 0
                    current_row_spans = []
                    estimated_height = 0

                    for cell_index, cell_def in enumerate(row_template.get("cells", [])):
                        c_span = cell_def.get("colspan", 1)
                        start_c = col_cursor
                        end_c = col_cursor + c_span - 1
                        col_cursor += c_span
                        current_row_spans.append((start_c, end_c))

                        if end_c > start_c:
                            t.cell(r, start_c).merge(t.cell(r, end_c))

                        cell_text = cell_texts[cell_index]

                        align = cell_def.get("align", "LEFT")
                        bold = cell_def.get("bold", False)
                        size_pt = cell_def.get("size_pt", 12.0)
                        self.sm.format_cell(
                            t.cell(r, start_c),
                            cell_text,
                            align=align,
                            bold=bold,
                            size_pt=size_pt,
                            style_override=lesson_override
                        )
                        if sec.get("paragraph_groups"):
                            cell = t.cell(r, start_c)
                            # The atomic row already keeps heading + body intact.
                            # Paragraph keepNext inside such a row can make Word
                            # pull the following row along, wasting page capacity.
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.keep_with_next = False
                            tc_pr = cell._tc.get_or_add_tcPr()
                            for node in tc_pr.xpath("w:vAlign"):
                                tc_pr.remove(node)
                            tc_pr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="top"/>'))
                            hidden_edges = []
                            if continuous_borders:
                                group_edges[r] = (["top"] if block_index == 0 else []) + (["bottom"] if block_index == block_count - 1 else [])
                            else:
                                if block_index:
                                    hidden_edges.append('<w:top w:val="nil"/>')
                                if block_index < block_count - 1:
                                    hidden_edges.append('<w:bottom w:val="nil"/>')
                            if hidden_edges:
                                tc_pr.insert_element_before(parse_xml(f'<w:tcBorders {nsdecls("w")}>{"".join(hidden_edges)}</w:tcBorders>'), "w:shd", "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign")
                            tc_pr.insert_element_before(parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="0" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/></w:tcMar>'), "w:textDirection", "w:tcFitText", "w:vAlign")
                            # A conservative guard, not a page-count predictor.
                            # Oversized rows must remain splittable: Word can
                            # clip an over-page row marked cantSplit.
                            margins = tbl_conf.get("cell_margins_dxa", {})
                            horizontal_margin = max(22, (margins.get("left", 108) + margins.get("right", 108)) / 20)
                            width_pt = sum(grid_cols_dxa[start_c:end_c + 1]) / 20 - horizontal_margin
                            cell_height = 0
                            for paragraph in cell.paragraphs:
                                font_size = max((run.font.size.pt for run in paragraph.runs if run.font.size), default=12)
                                chars_per_line = max(1, int(width_pt / font_size))
                                paragraph_format = paragraph.paragraph_format
                                spacing = paragraph_format.line_spacing
                                line_height = (
                                    max(font_size * 1.5, spacing.pt)
                                    if hasattr(spacing, "pt")
                                    else font_size * max(1.5, spacing or 1.5)
                                )
                                cell_height += (math.ceil(len(paragraph.text) / chars_per_line) + 1) * line_height
                                cell_height += sum(value.pt for value in (paragraph_format.space_before, paragraph_format.space_after) if value is not None)
                            estimated_height = max(estimated_height, cell_height)

                    row_spans.append(current_row_spans)
                    if sec.get("paragraph_groups"):
                        section = doc.sections[-1]
                        usable_height = (section.page_height - section.top_margin - section.bottom_margin) / 12700
                        max_unsplit_fraction = sec.get("max_unsplit_height_fraction", 0.82)
                        if estimated_height < usable_height * max_unsplit_fraction:
                            structural_rows.append(r)
                        else:
                            oversized_rows.append(r)
                            warnings.warn("教学过程包含较长的段落组，保留跨页能力；须在 Word/WPS 中核验该段分页。", stacklevel=2)

            elif sec_type == "reflection_blank":
                r = current_row_idx
                current_row_idx += 1
                col_cursor = 0
                current_row_spans = []

                def_lines = sec.get("lesson_0_lines", 4) if idx == 0 else sec.get("default_lines", 6)
                refl_lines = lesson.get("reflection_lines", lesson_override.get("reflection_lines", def_lines))
                # N requested lines should produce N lines, not N+1. Keeping the
                # default region compact prevents an otherwise empty reflection
                # box from being pushed onto its own page.
                blank_text = "\n" * max(0, int(refl_lines) - 1)

                for cell_def in sec.get("cells", [{"colspan": len(grid_cols_dxa), "align": "LEFT"}]):
                    c_span = cell_def.get("colspan", len(grid_cols_dxa))
                    start_c = col_cursor
                    end_c = col_cursor + c_span - 1
                    col_cursor += c_span
                    current_row_spans.append((start_c, end_c))

                    if end_c > start_c:
                        t.cell(r, start_c).merge(t.cell(r, end_c))

                    align = cell_def.get("align", "LEFT")
                    bold = cell_def.get("bold", False)
                    size_pt = cell_def.get("size_pt", 12.0)
                    self.sm.format_cell(
                        t.cell(r, start_c),
                        blank_text,
                        align=align,
                        bold=bold,
                        size_pt=size_pt,
                        style_override=lesson_override
                    )

                row_spans.append(current_row_spans)
                structural_rows.append(r)

        # Apply exact calculated widths
        cant_split = tbl_conf.get("cant_split", False)
        self.sm.apply_dynamic_cell_widths(t, grid_cols_dxa, row_spans_list=row_spans, cant_split=cant_split)
        for row_idx in repeating_header_rows:
            self.sm.set_repeat_table_header(t.rows[row_idx])
            for cell in t.rows[row_idx].cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
        for row_idx in structural_rows:
            self.sm.set_row_cant_split(t.rows[row_idx])
        # Safety takes precedence over a template-wide no-split preference.
        for row_idx in oversized_rows:
            row = t.rows[row_idx]
            for node in row._tr.xpath("w:trPr/w:cantSplit"):
                node.getparent().remove(node)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_together = False
                    paragraph.paragraph_format.keep_with_next = False
        if continuous_borders:
            for row_index, row in enumerate(t.rows):
                edges = group_edges.get(row_index, ["top", "bottom"])
                for tc in row._tr.tc_lst:
                    if edges:
                        borders = "".join(f'<w:{edge} w:val="single" w:color="000000" w:sz="4"/>' for edge in edges)
                        tc.get_or_add_tcPr().insert_element_before(parse_xml(f'<w:tcBorders {nsdecls("w")}>{borders}</w:tcBorders>'), "w:shd", "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign")
        keep_reflection = self.tm.get_section("pagination", default={}).get("keep_reflection_with_previous", False)
        keep_before_rows = [r for r in reflection_starts if r - 1 in structural_rows and r - 1 not in oversized_rows] if keep_reflection else []
        self._split_table_at_rows(t, split_rows, continuous_process_borders=continuous_borders, keep_before_rows=keep_before_rows)

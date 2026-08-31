"""OOXML regression tests for paragraph-group teaching-process rows."""

import copy
import json
import sys
import unittest
import warnings
from pathlib import Path


SKILL_DIR = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(SKILL_DIR / "scripts"))

from docx.table import _Cell  # noqa: E402
from doc_builder import TeachingPlanDocBuilder  # noqa: E402
from data_normalizer import DataNormalizer  # noqa: E402
from pagination import split_paragraph_blocks  # noqa: E402
from style_manager import StyleManager  # noqa: E402


def template_config():
    return json.loads((SKILL_DIR / "templates" / "default_vocational.json").read_text(encoding="utf-8"))


def process_table(document):
    matches = [
        table for table in document.tables
        if table.rows and table.rows[0].cells[0].text.replace(" ", "") == "教学过程"
    ]
    if len(matches) != 1:
        raise AssertionError(f"Expected one process table, got {len(matches)}")
    return matches[0]


def physical_cell_texts(row, table):
    # row.cells repeats the same design cell for its two grid columns.
    return [_Cell(tc, table).text for tc in row._tr.tc_lst]


def short_lesson():
    return {
        "lesson_num": "第1、2课时",
        "title": "分页回归测试",
        "stages": [
            {
                "stage_name": "任务初探",
                "content": "\n活动1\n\n教师活动\n1. 唯一观察指令甲\n2. 唯一比较指令乙\n学生活动\n1. 唯一记录指令丙\n",
                "purpose": "1. 唯一设计意图丁\n2. 唯一设计意图戊",
            },
            {
                "stage_index": "（自定）",
                "stage_name": "任务评价",
                "design": "多元评价\n1. 唯一评价指令己\n2. 唯一反馈指令庚",
                "intent": "唯一设计意图辛",
            },
        ],
    }


class ProcessPaginationTests(unittest.TestCase):
    def build(self, lesson, template=None):
        return TeachingPlanDocBuilder({"lessons": [lesson]}, template=template).build_document()

    def assert_grouped_text_is_lossless(self, lesson):
        original = copy.deepcopy(lesson)
        normalized = DataNormalizer.normalize_lesson(lesson)
        unbalanced = template_config()
        for section in unbalanced["lesson_table"]["sections"]:
            section["balance_columns"] = False
        table = process_table(self.build(lesson, unbalanced))
        body_rows = table.rows[2:]
        cursor = 0
        heading_check = StyleManager().is_subhead
        for stage in normalized["stages"]:
            label = f"{stage['stage_index']}\n{stage['stage_name']}"
            design_blocks = split_paragraph_blocks(stage["design"], heading_check)
            intent_blocks = split_paragraph_blocks(stage["intent"], heading_check)
            count = max(1, len(design_blocks), len(intent_blocks))
            actual = [physical_cell_texts(row, table) for row in body_rows[cursor:cursor + count]]
            self.assertEqual(len(actual), count)
            for index, cells in enumerate(actual):
                self.assertEqual(cells, [
                    label if index == 0 else "",
                    design_blocks[index] if index < len(design_blocks) else "",
                    intent_blocks[index] if index < len(intent_blocks) else "",
                ])
            # Reconstruct only real blocks; blank padding belongs to the other
            # column and must not masquerade as duplicated source newlines.
            self.assertEqual("\n".join(cells[1] for cells in actual[:len(design_blocks)]), stage["design"])
            self.assertEqual("\n".join(cells[2] for cells in actual[:len(intent_blocks)]), stage["intent"])
            self.assertEqual([cells[0] for cells in actual if cells[0]], [label])
            cursor += count
        self.assertEqual(cursor, len(body_rows))
        self.assertEqual(lesson, original, "Builder must not mutate source content")

    def test_normalized_design_label_and_intent_are_complete_once(self):
        self.assert_grouped_text_is_lossless(short_lesson())

    def test_legacy_stages_preserve_normalized_text(self):
        lesson = {
            "lesson_num": "第3、4课时",
            "title": "旧格式分页回归",
            "stage_intro": "教师活动\n1. 旧格式独有内容甲\n学生活动\n1. 旧格式独有内容乙",
            "intent_intro": "旧格式独有意图甲\n旧格式独有意图乙",
            "stage_execute": "活动1\n教师活动\n1. 旧格式独有内容丙",
            "intent_execute": "旧格式独有意图丙",
        }
        self.assert_grouped_text_is_lossless(lesson)

    def test_real_example_stages_preserve_normalized_text(self):
        for path in sorted((SKILL_DIR / "examples").glob("*.json")):
            with self.subTest(example=path.name):
                content = json.loads(path.read_text(encoding="utf-8"))
                self.assert_grouped_text_is_lossless(content["lessons"][0])

    def test_default_short_groups_are_unsplittable_physical_rows(self):
        table = process_table(self.build(short_lesson()))
        self.assertGreater(len(table.rows[2:]), len(short_lesson()["stages"]))
        for row in table.rows[2:]:
            self.assertTrue(row._tr.xpath("w:trPr/w:cantSplit"))
            self.assertFalse(row._tr.xpath("w:tc/w:tcPr/w:vMerge"))
            self.assertEqual(len(row._tr.tc_lst), 3)

    def test_both_header_rows_repeat_and_keep_with_next(self):
        table = process_table(self.build(short_lesson()))
        for row in table.rows[:2]:
            self.assertTrue(row._tr.xpath("w:trPr/w:tblHeader"))
            self.assertTrue(row._tr.xpath("w:trPr/w:cantSplit"))
            for tc in row._tr.tc_lst:
                for paragraph in _Cell(tc, table).paragraphs:
                    self.assertIs(paragraph.paragraph_format.keep_with_next, True)

    def test_explicit_false_retains_one_row_per_stage(self):
        template = template_config()
        for section in template["lesson_table"]["sections"]:
            if section.get("type") == "repeater":
                section["paragraph_groups"] = False
        lesson = short_lesson()
        stages = DataNormalizer.normalize_lesson(lesson)["stages"]
        table = process_table(self.build(lesson, template))
        body_rows = table.rows[2:]
        self.assertEqual(len(body_rows), len(stages))
        for row, stage in zip(body_rows, stages):
            self.assertEqual(physical_cell_texts(row, table), [
                f"{stage['stage_index']}\n{stage['stage_name']}", stage["design"], stage["intent"],
            ])
            self.assertFalse(row._tr.xpath("w:trPr/w:cantSplit"))
            self.assertFalse(row._tr.xpath("w:tc/w:tcPr/w:tcBorders"))

    def test_oversized_paragraph_warns_and_remains_splittable(self):
        long_text = "超长段落全文保留" * 1000
        lesson = {
            "lesson_num": "第5、6课时",
            "title": "超长段落回归",
            "stages": [{"stage_name": "任务执行", "design": f"教师活动\n{long_text}\n短段落尾标", "intent": "独有意图"}],
        }
        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            table = process_table(self.build(lesson))
        pagination_warnings = [warning for warning in caught if "较长的段落组" in str(warning.message)]
        self.assertEqual(len(pagination_warnings), 1)
        matching_rows = [row for row in table.rows[2:] if long_text in physical_cell_texts(row, table)[1]]
        self.assertEqual(len(matching_rows), 1)
        self.assertFalse(matching_rows[0]._tr.xpath("w:trPr/w:cantSplit"))
        self.assertEqual(physical_cell_texts(matching_rows[0], table)[1], f"教师活动\n{long_text}")
        short_rows = [row for row in table.rows[2:] if physical_cell_texts(row, table)[1] == "短段落尾标"]
        self.assertEqual(len(short_rows), 1)
        self.assertTrue(short_rows[0]._tr.xpath("w:trPr/w:cantSplit"))

    def test_oversized_safety_overrides_global_cant_split(self):
        template = template_config()
        template["lesson_table"]["table"]["cant_split"] = True
        long_text = "全局不可拆行配置不能覆盖超长组安全回退" * 1000
        lesson = {
            "lesson_num": "第7、8课时",
            "title": "全局分页配置回归",
            "stages": [{"stage_name": "任务执行", "design": f"教师活动\n{long_text}\n短组仍受保护", "intent": "安全优先"}],
        }
        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            table = process_table(self.build(lesson, template))
        self.assertEqual(sum("较长的段落组" in str(item.message) for item in caught), 1)
        long_rows = [row for row in table.rows[2:] if long_text in physical_cell_texts(row, table)[1]]
        self.assertEqual(len(long_rows), 1)
        self.assertFalse(long_rows[0]._tr.xpath("w:trPr/w:cantSplit"))
        self.assertEqual(physical_cell_texts(long_rows[0], table)[1], f"教师活动\n{long_text}")
        for tc in long_rows[0]._tr.tc_lst:
            for paragraph in _Cell(tc, table).paragraphs:
                self.assertIs(paragraph.paragraph_format.keep_together, False)
                self.assertIs(paragraph.paragraph_format.keep_with_next, False)
        short_rows = [row for row in table.rows[2:] if physical_cell_texts(row, table)[1] == "短组仍受保护"]
        self.assertEqual(len(short_rows), 1)
        self.assertTrue(short_rows[0]._tr.xpath("w:trPr/w:cantSplit"))

    def test_extreme_custom_spacing_uses_splittable_fallback(self):
        lesson = {
            "lesson_num": "第9课时",
            "title": "自定义行距回归",
            "style_override": {"line_spacing": 50},
            "stages": [{"stage_name": "任务执行", "design": "教师活动\n保留完整短正文", "intent": "检查安全回退"}],
        }
        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            table = process_table(self.build(lesson))
        self.assertTrue(caught)
        self.assertFalse(table.rows[2]._tr.xpath("w:trPr/w:cantSplit"))


if __name__ == "__main__":
    unittest.main()

# -*- coding: utf-8 -*-

import json
import sys
import tempfile
import unittest
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path


SKILL_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = SKILL_DIR / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))

import docx  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from font_preflight import build_substitution_override  # noqa: E402
from render import ContentValidationError, generate_document_from_file, generate_single_document  # noqa: E402
from doc_builder import TeachingPlanDocBuilder  # noqa: E402
from template_manager import TemplateManager  # noqa: E402
from style_manager import StyleManager  # noqa: E402


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


class RenderQualityTests(unittest.TestCase):
    def test_cell_paragraphs_do_not_orphan_headings_or_link_rows(self):
        manager = StyleManager()
        cell = docx.Document().add_table(rows=1, cols=1).cell(0, 0)
        manager.format_cell(cell, "活动3 检查作品\n教师活动\n说明要求。\n学生活动\n完成操作。")
        self.assertEqual(
            [p.paragraph_format.keep_with_next for p in cell.paragraphs],
            [True, True, False, True, False],
        )
        for paragraph in cell.paragraphs:
            self.assertTrue(paragraph.paragraph_format.keep_together)
            self.assertTrue(paragraph.paragraph_format.widow_control)
        self.assertFalse(cell._tc.getparent().xpath("w:trPr/w:cantSplit"))

        manager.format_cell(cell, "（二）\n分析任务", bold=True)
        self.assertTrue(cell.paragraphs[0].paragraph_format.keep_with_next)
        self.assertFalse(cell.paragraphs[-1].paragraph_format.keep_with_next)

    def test_invalid_content_is_blocked_at_render_boundary(self):
        with self.assertRaises(ContentValidationError):
            generate_single_document({}, quiet=True)

    def test_single_lesson_profile_renders_without_cover_or_schedule(self):
        example = json.loads(next((SKILL_DIR / "examples").glob("*.json")).read_text(encoding="utf-8"))
        content = {
            "document_profile": "single_lesson",
            "document_status": "final",
            "output_filename": "单课时教案.docx",
            "lessons": [example["lessons"][0]],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            output, document = generate_single_document(content, output_dir=temp_dir, quiet=True)
            self.assertTrue(Path(output).exists())
            all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertNotIn("教学进度表", all_text)
            self.assertEqual(len(document.tables), 3)

    def test_default_template_does_not_supply_school_term_or_date_facts(self):
        template = TemplateManager()
        cover = template.get_section("cover")
        for field in ("subtitle", "school", "date"):
            with self.subTest(field=field):
                expression = cover[field]["bind"]
                self.assertEqual(template.evaluate_expression(expression, {"cover": {}}), "")

    def test_font_override_is_explicit_and_applied_to_runs(self):
        report = {
            "font_checks": [
                {
                    "requested": "Missing Chinese Font",
                    "suggested_fallback": "Installed Chinese Font",
                }
            ]
        }
        override = build_substitution_override(report)
        self.assertEqual(
            override,
            {"fonts": {"substitutions": {"Missing Chinese Font": "Installed Chinese Font"}}},
        )
        manager = StyleManager(
            styles_config={
                "fonts": {
                    "default_font": "Missing Chinese Font",
                    "ascii_font": "Times New Roman",
                    "substitutions": {"Missing Chinese Font": "Installed Chinese Font"},
                }
            }
        )
        paragraph = docx.Document().add_paragraph()
        run = manager.add_run(paragraph, "教案", font_name="Missing Chinese Font")
        fonts = run._r.get_or_add_rPr().find(qn("w:rFonts"))
        self.assertEqual(fonts.get(qn("w:eastAsia")), "Installed Chinese Font")
        self.assertEqual(fonts.get(qn("w:ascii")), "Installed Chinese Font")

    def test_bundled_examples_have_deterministic_table_geometry(self):
        for example in sorted((SKILL_DIR / "examples").glob("*.json")):
            with self.subTest(example=example.name), tempfile.TemporaryDirectory() as temp_dir:
                output, _doc = generate_document_from_file(
                    str(example), output_dir=temp_dir, quiet=True
                )
                with zipfile.ZipFile(output) as archive:
                    root = ET.fromstring(archive.read("word/document.xml"))

                self.assertFalse(root.findall(f".//{W}br[@{W}type='page']"))
                self.assertTrue(root.findall(f".//{W}tblHeader"))
                tables = root.findall(f".//{W}tbl")
                self.assertGreater(len(tables), 2)
                for table_index, table in enumerate(tables):
                    width_node = table.find(f"./{W}tblPr/{W}tblW")
                    self.assertIsNotNone(width_node, table_index)
                    self.assertEqual(width_node.get(f"{W}type"), "dxa")
                    table_width = int(width_node.get(f"{W}w"))
                    grid_width = sum(
                        int(column.get(f"{W}w"))
                        for column in table.findall(f"./{W}tblGrid/{W}gridCol")
                    )
                    self.assertEqual(table_width, grid_width, table_index)
                    for row_index, row in enumerate(table.findall(f"./{W}tr")):
                        cell_width = sum(
                            int(node.get(f"{W}w"))
                            for node in row.findall(f"./{W}tc/{W}tcPr/{W}tcW")
                        )
                        self.assertEqual(
                            cell_width,
                            table_width,
                            f"table {table_index}, row {row_index}",
                        )

    def test_lesson_fragments_stay_separate_in_word(self):
        example = next((SKILL_DIR / "examples").glob("*.json"))
        with tempfile.TemporaryDirectory() as temp_dir:
            output, _doc = generate_document_from_file(
                str(example), output_dir=temp_dir, quiet=True
            )
            with zipfile.ZipFile(output) as archive:
                root = ET.fromstring(archive.read("word/document.xml"))

        body = root.find(f".{W}body")
        children = list(body)
        separators = []
        for index in range(1, len(children) - 1):
            if (
                children[index - 1].tag == f"{W}tbl"
                and children[index].tag == f"{W}p"
                and children[index + 1].tag == f"{W}tbl"
            ):
                spacing = children[index].find(f"./{W}pPr/{W}spacing")
                if spacing is not None and spacing.get(f"{W}line") == "20":
                    separators.append(children[index])
        self.assertGreater(len(separators), 0)

        process_tables = []
        for table in body.findall(f"./{W}tbl"):
            rows = table.findall(f"./{W}tr")
            first_text = "".join(rows[0].itertext()) if rows else ""
            if "教  学  过  程" in first_text:
                process_tables.append(rows)
        self.assertGreater(len(process_tables), 0)
        for rows in process_tables:
            for row in rows[:2]:
                row_properties = list(row.find(f"./{W}trPr"))
                cant_split_index = next(
                    index for index, node in enumerate(row_properties) if node.tag == f"{W}cantSplit"
                )
                header_index = next(
                    index for index, node in enumerate(row_properties) if node.tag == f"{W}tblHeader"
                )
                self.assertLess(cant_split_index, header_index)

    def test_lesson_start_pagination_is_configurable(self):
        example = next((SKILL_DIR / "examples").glob("*.json"))
        content = json.loads(example.read_text(encoding="utf-8"))
        template_data = json.loads(
            (SKILL_DIR / "templates" / "default_vocational.json").read_text(encoding="utf-8")
        )

        counts = {}
        for mode in ("new_page", "flow"):
            template_data["pagination"]["lesson_start"] = mode
            document = TeachingPlanDocBuilder(
                content,
                template=TemplateManager(template_data),
            ).build_document()
            root = document._element
            counts[mode] = len(root.findall(f".//{W}pageBreakBefore"))

        lesson_count = len(content["lessons"])
        self.assertEqual(1 + lesson_count, counts["new_page"])
        self.assertEqual(2, counts["flow"])


if __name__ == "__main__":
    unittest.main()

"""Independent packing invariants and default-template OOXML integration checks."""

import copy
import itertools
import json
import sys
import unittest
import warnings
from pathlib import Path

from docx.oxml.ns import qn
from docx.table import _Cell

SKILL_DIR = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(SKILL_DIR / "scripts"))

from data_normalizer import DataNormalizer
from doc_builder import TeachingPlanDocBuilder
from pagination import balance_cell_blocks, estimate_text_height, split_paragraph_blocks
from style_manager import StyleManager


def default_template():
    return json.loads((SKILL_DIR / "templates" / "default_vocational.json").read_text(encoding="utf-8"))


def build_process(lesson, template=None):
    document = TeachingPlanDocBuilder({"lessons": [lesson]}, template=template).build_document()
    matches = [table for table in document.tables if table.rows[0].cells[0].text.replace(" ", "") == "教学过程"]
    if len(matches) != 1:
        raise AssertionError(f"Expected one process table; got {len(matches)}")
    return matches[0]


def build_with_process_reflection(lesson, template=None):
    document = TeachingPlanDocBuilder({"lessons": [lesson]}, template=template).build_document()
    process_matches = [
        table for table in document.tables
        if table.rows and table.rows[0].cells[0].text.replace(" ", "") == "教学过程"
    ]
    reflection_matches = [
        table for table in document.tables
        if table.rows and table.rows[0].cells[0].text.replace(" ", "") == "教学反思"
    ]
    if len(process_matches) != 1 or len(reflection_matches) != 1:
        raise AssertionError(
            f"Expected one process and one reflection table; got "
            f"{len(process_matches)} and {len(reflection_matches)}"
        )
    process = process_matches[0]
    reflection = reflection_matches[0]
    process_index = document._element.body.index(process._tbl)
    reflection_index = document._element.body.index(reflection._tbl)
    if reflection_index != process_index + 2:
        raise AssertionError("Expected exactly one separator paragraph before reflection")
    separator = document._element.body[process_index + 1]
    if not separator.tag.endswith("}p"):
        raise AssertionError("Expected a paragraph separator before reflection")
    return document, process, separator, reflection


def physical_texts(table):
    return [[_Cell(tc, table).text for tc in row._tr.tc_lst] for row in table.rows[2:]]


class BalancedPackingTests(unittest.TestCase):
    def test_all_atomic_blocks_preserved_once_in_order_including_empty_blocks(self):
        # [] is padding, whereas [""] is one genuine source block.
        choices = ([], [""], ["\n"], ["a", ""], ["", "b", "\n"], ["标题\n正文", "尾段\n"])
        for columns in itertools.product(choices, repeat=3):
            before = copy.deepcopy(columns)
            rows = balance_cell_blocks(columns, lambda _, text: len(text), max_height=8)
            for index, blocks in enumerate(columns):
                self.assertEqual([block for row in rows for block in row[index]], blocks)
            self.assertEqual(columns, before)

    def test_short_blocks_fill_tall_neighbor_without_reordering(self):
        columns = [["abcdefgh"], ["a", "b", "c", "d"]]
        rows = balance_cell_blocks(columns, lambda _, text: len(text))
        self.assertEqual(rows, [[["abcdefgh"], ["a", "b", "c", "d"]]])

    def test_already_balanced_rows_are_not_merged(self):
        columns = [["aaaa", "bbbb"], ["cccc", "dddd"]]
        rows = balance_cell_blocks(columns, lambda _, text: len(text))
        self.assertEqual(rows, [[["aaaa"], ["cccc"]], [["bbbb"], ["dddd"]]])

    def test_oversized_first_block_does_not_absorb_later_blocks_from_any_column(self):
        columns = [["x" * 200, "尾段"], ["意图甲", "意图乙", "意图丙"]]
        rows = balance_cell_blocks(columns, lambda _, text: len(text), max_height=160)
        self.assertEqual(rows[0], [["x" * 200], ["意图甲"]])
        self.assertEqual([block for row in rows[1:] for block in row[1]], ["意图乙", "意图丙"])

    def test_height_estimate_tracks_width_font_spacing_and_blank_lines(self):
        text = "中文排版" * 8
        baseline = estimate_text_height(text, 120)
        self.assertGreater(estimate_text_height(text, 60), baseline)
        self.assertGreater(estimate_text_height(text, 120, size_pt=16), baseline)
        self.assertGreater(estimate_text_height(text, 120, line_spacing=2), baseline)
        self.assertGreater(estimate_text_height(text + "\n\n", 120), baseline)
        self.assertGreater(estimate_text_height(text, 120, paragraph_spacing=10), baseline)


class BalancedLayoutIntegrationTests(unittest.TestCase):
    def assert_column_partition(self, source, row_texts):
        """Each output cell must contain consecutive complete source blocks.

        Do not strip/filter blank strings: a final empty source block and a
        genuinely padded cell are distinguished using source-block consumption.
        """
        blocks = split_paragraph_blocks(source, StyleManager().is_subhead)
        cursor = 0
        recovered = []
        for text in row_texts:
            if cursor == len(blocks):
                self.assertEqual(text, "", "Padding must not invent text")
                continue
            match = next((end for end in range(cursor + 1, len(blocks) + 1)
                          if "\n".join(blocks[cursor:end]) == text), None)
            self.assertIsNotNone(match, f"Output cell is not an intact block sequence: {text!r}")
            recovered.extend(blocks[cursor:match])
            cursor = match
        self.assertEqual(cursor, len(blocks), "Source blocks were dropped")
        self.assertEqual("\n".join(recovered), source)

    def test_default_balancing_preserves_both_columns_and_stage_labels(self):
        lesson = {
            "lesson_num": "第1课时", "title": "双列平衡测试",
            "stages": [
                {"stage_name": "任务初探", "design": "\n活动1\n\n教师活动\n" + "观察并记录。" * 12 + "\n学生活动\n操作甲\n操作乙\n\n",
                 "intent": "\n意图甲\n意图乙\n意图丙\n"},
                {"stage_name": "任务评价", "design": "反馈甲\n反馈乙\n", "intent": "评价甲\n\n评价乙\n\n"},
            ],
        }
        before = copy.deepcopy(lesson)
        stages = DataNormalizer.normalize_lesson(lesson)["stages"]
        table = build_process(lesson)
        rows = physical_texts(table)
        starts = [index for index, cells in enumerate(rows) if cells[0]] + [len(rows)]
        self.assertEqual(len(starts) - 1, len(stages))
        for stage, start, end in zip(stages, starts, starts[1:]):
            group = rows[start:end]
            self.assertEqual(group[0][0], f"{stage['stage_index']}\n{stage['stage_name']}")
            self.assertTrue(all(cells[0] == "" for cells in group[1:]))
            self.assert_column_partition(stage["design"], [cells[1] for cells in group])
            self.assert_column_partition(stage["intent"], [cells[2] for cells in group])
        self.assertEqual(lesson, before)

    def test_default_packing_reduces_row_count_and_keeps_short_rows_protected(self):
        lesson = {"lesson_num": "第2课时", "title": "减少局部留白",
                  "stages": [{"stage_name": "任务执行", "design": "较长正文" * 22,
                              "intent": "意图甲\n意图乙\n意图丙\n意图丁"}]}
        old_template = default_template()
        for section in old_template["lesson_table"]["sections"]:
            section["balance_columns"] = False
        old_table = build_process(lesson, old_template)
        table = build_process(lesson)
        self.assertLess(len(table.rows), len(old_table.rows))
        for row in table.rows[2:]:
            self.assertTrue(row._tr.xpath("w:trPr/w:cantSplit"))
            self.assertFalse(row._tr.xpath("w:trPr/w:trHeight[@w:hRule='exact']"))

    def test_grouped_heading_body_stays_in_one_row_without_paragraph_links_between_rows(self):
        lesson = {"lesson_num": "第2课时", "title": "标题和正文不跨行串联",
                  "stages": [{"stage_name": "任务执行",
                              "design": "活动1\n教师活动\n唯一观察正文甲\n唯一追加正文乙\n学生活动\n唯一操作正文丙",
                              "intent": "唯一意图甲\n唯一意图乙\n唯一意图丙"}]}
        table = build_process(lesson)
        text_rows = physical_texts(table)
        self.assertTrue(any("活动1\n教师活动\n唯一观察正文甲" in cells[1] for cells in text_rows))
        self.assertTrue(any("学生活动\n唯一操作正文丙" in cells[1] for cells in text_rows))
        self.assertGreater(len(table.rows[2:]), 1)
        for row_index, row in enumerate(table.rows[2:]):
            self.assertTrue(row._tr.xpath("w:trPr/w:cantSplit"))
            for tc in row._tr.tc_lst:
                for paragraph in _Cell(tc, table).paragraphs:
                    # The final short row is intentionally anchored to the
                    # following reflection heading; earlier grouped rows must
                    # not form a page-wide keepNext chain.
                    self.assertIs(paragraph.paragraph_format.keep_with_next, row_index == len(table.rows[2:]) - 1)
                    self.assertIs(paragraph.paragraph_format.keep_together, True)
        # Header repetition is a separate mechanism and must not be weakened.
        for row in table.rows[:2]:
            self.assertTrue(row._tr.xpath("w:trPr/w:tblHeader"))
            for tc in row._tr.tc_lst:
                for paragraph in _Cell(tc, table).paragraphs:
                    self.assertIs(paragraph.paragraph_format.keep_with_next, True)

    def test_default_reflection_anchor_links_only_final_short_process_row_and_separator(self):
        lesson = {
            "lesson_num": "第2课时", "title": "反思尾锚默认开启",
            "stages": [{
                "stage_name": "任务执行",
                "design": "活动1\n教师活动\n观察甲\n学生活动\n操作乙",
                "intent": "意图甲\n意图乙",
            }],
        }
        _, process, separator, reflection = build_with_process_reflection(lesson)
        body_rows = process.rows[2:]
        self.assertGreater(len(body_rows), 1)

        for row in body_rows[:-1]:
            for tc in row._tr.tc_lst:
                for paragraph in _Cell(tc, process).paragraphs:
                    self.assertIs(paragraph.paragraph_format.keep_with_next, False)
        for tc in body_rows[-1]._tr.tc_lst:
            for paragraph in _Cell(tc, process).paragraphs:
                self.assertIs(paragraph.paragraph_format.keep_with_next, True)

        self.assertTrue(separator.xpath("w:pPr/w:keepNext"))
        for paragraph in reflection.rows[0].cells[0].paragraphs:
            self.assertIs(paragraph.paragraph_format.keep_with_next, True)
        for paragraph in reflection.rows[-1].cells[0].paragraphs:
            self.assertIs(paragraph.paragraph_format.keep_with_next, False)

    def test_disabled_reflection_anchor_restores_unlinked_process_tail_and_separator(self):
        template = default_template()
        template["pagination"]["keep_reflection_with_previous"] = False
        lesson = {
            "lesson_num": "第2课时", "title": "反思尾锚关闭",
            "stages": [{
                "stage_name": "任务执行",
                "design": "活动1\n教师活动\n观察甲\n学生活动\n操作乙",
                "intent": "意图甲\n意图乙",
            }],
        }
        _, process, separator, reflection = build_with_process_reflection(lesson, template)
        for row in process.rows[2:]:
            for tc in row._tr.tc_lst:
                for paragraph in _Cell(tc, process).paragraphs:
                    self.assertIs(paragraph.paragraph_format.keep_with_next, False)
        self.assertFalse(separator.xpath("w:pPr/w:keepNext"))

        # The reflection heading still owns its blank writing area. Disabling
        # the cross-table tail anchor must not break that independent link.
        for paragraph in reflection.rows[0].cells[0].paragraphs:
            self.assertIs(paragraph.paragraph_format.keep_with_next, True)

    def test_oversized_final_process_row_is_never_forced_into_reflection_anchor(self):
        long_text = "超长尾行不得锚定反思" * 1000
        lesson = {
            "lesson_num": "第3课时", "title": "超长尾行安全回退",
            "stages": [{"stage_name": "任务执行", "design": long_text, "intent": "单一意图"}],
        }
        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            _, process, separator, _ = build_with_process_reflection(lesson)
        self.assertTrue(any("较长的段落组" in str(item.message) for item in caught))
        self.assertEqual(len(process.rows[2:]), 1)
        final_row = process.rows[-1]
        self.assertFalse(final_row._tr.xpath("w:trPr/w:cantSplit"))
        for tc in final_row._tr.tc_lst:
            for paragraph in _Cell(tc, process).paragraphs:
                self.assertIs(paragraph.paragraph_format.keep_with_next, False)
                self.assertIs(paragraph.paragraph_format.keep_together, False)
        self.assertFalse(separator.xpath("w:pPr/w:keepNext"))

    def test_oversized_row_remains_splittable_while_following_short_rows_are_protected(self):
        long_text = "超长原子段落" * 1000
        lesson = {"lesson_num": "第3课时", "title": "超长行安全回退",
                  "stages": [{"stage_name": "任务执行", "design": long_text + "\n后续短段落",
                              "intent": "意图甲\n意图乙\n意图丙"}]}
        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            table = build_process(lesson)
        rows = physical_texts(table)
        self.assertEqual(rows[0][1:], [long_text, "意图甲"])
        self.assertFalse(table.rows[2]._tr.xpath("w:trPr/w:cantSplit"))
        self.assertTrue(any("较长的段落组" in str(item.message) for item in caught))
        self.assertTrue(any("后续短段落" == cells[1] for cells in rows[1:]))
        for row in table.rows[3:]:
            self.assertTrue(row._tr.xpath("w:trPr/w:cantSplit"))

    def test_continuous_borders_preserve_outer_frame_without_cell_nil_overrides(self):
        lesson = {"lesson_num": "第4课时", "title": "连续外框",
                  "stages": [{"stage_name": "任务执行", "design": "正文甲\n正文乙\n正文丙", "intent": "意图甲\n意图乙\n意图丙"}]}
        table = build_process(lesson)
        self.assertEqual(len(table._tbl.xpath("w:tblPr/w:tblBorders")), 1)
        for edge in ("top", "bottom", "left", "right", "insideV"):
            nodes = table._tbl.xpath(f"w:tblPr/w:tblBorders/w:{edge}")
            self.assertEqual(len(nodes), 1)
            self.assertNotIn(nodes[0].get(qn("w:val")), ("none", "nil"))
        self.assertEqual(table._tbl.xpath("w:tblPr/w:tblBorders/w:insideH")[0].get(qn("w:val")), "nil")
        self.assertFalse(table._tbl.xpath("w:tr/w:tc/w:tcPr/w:tcBorders/*[@w:val='nil' or @w:val='none']"))
        for row in (table.rows[2], table.rows[-1]):
            for tc in row._tr.tc_lst:
                self.assertLessEqual(len(tc.xpath("w:tcPr/w:tcBorders")), 1)
        for tc in table.rows[2]._tr.tc_lst:
            self.assertTrue(tc.xpath("w:tcPr/w:tcBorders/w:top[@w:val='single']"))
        for tc in table.rows[-1]._tr.tc_lst:
            self.assertTrue(tc.xpath("w:tcPr/w:tcBorders/w:bottom[@w:val='single']"))
        table_edge = table._tbl.xpath("w:tblPr/w:tblBorders/w:top")[0]
        for edge in table._tbl.xpath("w:tr/w:tc/w:tcPr/w:tcBorders/*"):
            for attribute in ("val", "color", "sz"):
                self.assertEqual(edge.get(qn("w:" + attribute)), table_edge.get(qn("w:" + attribute)))

    def test_missing_new_options_retains_legacy_unbalanced_grouped_layout(self):
        template = default_template()
        for section in template["lesson_table"]["sections"]:
            section.pop("balance_columns", None)
            section.pop("continuous_borders", None)
        lesson = {"lesson_num": "第5课时", "title": "旧模板兼容",
                  "stages": [{"stage_name": "任务执行", "design": "正文甲\n正文乙\n正文丙", "intent": "意图甲"}]}
        table = build_process(lesson, template)
        rows = physical_texts(table)
        self.assertEqual([cells[1:] for cells in rows], [["正文甲", "意图甲"], ["正文乙", ""], ["正文丙", ""]])
        self.assertTrue(table._tbl.xpath("w:tr/w:tc/w:tcPr/w:tcBorders/*[@w:val='nil']"))


if __name__ == "__main__":
    unittest.main()

import sys
import tempfile
import unittest
from pathlib import Path

import docx
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))
from style_manager import StyleManager
from technical_text import split_technical_text


def tokens(text):
    return [chunk for chunk, technical in split_technical_text(text) if technical]


class TechnicalTextTests(unittest.TestCase):
    def test_recognizes_commands_options_permissions_paths_and_ip(self):
        cases = {
            "讲解ls -l和chmod命令，查看/etc/passwd。": ["ls", "-l", "chmod", "/etc/passwd"],
            "记录rwxr-xr--与-rw-r--r--的区别": ["rwxr-xr--", "-rw-r--r--"],
            "路径~/work、../data、./run.sh和/var/log/": ["~/work", "../data", "./run.sh", "/var/log/"],
            "IP 192.168.10.100与192.168.10.0/24": ["192.168.10.100", "192.168.10.0/24"],
            "ls --all /tmp": ["ls", "--all", "/tmp"],
            "--help": ["--help"],
            "chmod 755 ./run.sh": ["chmod", "./run.sh"],
        }
        for text, expected in cases.items():
            with self.subTest(text=text):
                self.assertEqual(tokens(text), expected)

    def test_ordinary_english_numbers_and_invalid_ips_are_not_code(self):
        for text in (
            "Linux与Shell基础，共40课时，第1周，权限数字754与644。",
            "The cat and the head teacher touch the paper.",
            "学习cat单词，制作post-production短片。",
            "UID/GID、u/g/o、r/w/x和16/9。",
            "地址999.168.10.100和192.168.10.0/99无效。",
            "https://example.com/docs/path 与 user@example.com",
            "---------",
        ):
            with self.subTest(text=text):
                self.assertEqual(tokens(text), [])

    def test_recognition_preserves_all_source_characters(self):
        text = "  输入ls  -l\t/etc/passwd；权限-rwxr-xr--。\nIP 192.168.1.1，-3，中文路径/教学/课件。"
        self.assertEqual("".join(chunk for chunk, _ in split_technical_text(text)), text)

    def test_mixed_runs_preserve_cjk_font_and_body_size(self):
        document = docx.Document()
        cell = document.add_table(rows=1, cols=1).cell(0, 0)
        manager = StyleManager(style_override={"technical_text": {"font_name": "Consolas"}})
        source = "执行ls -l /etc/passwd，记录结果。"
        manager.format_cell(cell, source, font_name="楷体", size_pt=13.5)
        paragraph = cell.paragraphs[0]
        self.assertEqual(paragraph.text, source)
        for run in paragraph.runs:
            if not run.text:
                continue
            fonts = run._r.rPr.find(qn("w:rFonts"))
            self.assertEqual(fonts.get(qn("w:eastAsia")), "楷体")
            self.assertEqual(run.font.size.pt, 13.5)
            if run.text in ("ls", "-l", "/etc/passwd"):
                self.assertEqual(fonts.get(qn("w:ascii")), "Consolas")
                self.assertEqual(fonts.get(qn("w:hAnsi")), "Consolas")
            elif any("\u4e00" <= char <= "\u9fff" for char in run.text):
                self.assertEqual(fonts.get(qn("w:ascii")), "楷体")
        self.assertEqual(paragraph._p.pPr.find(qn("w:wordWrap")).get(qn("w:val")), "1")
        for tag in ("autoSpaceDE", "autoSpaceDN", "snapToGrid"):
            self.assertEqual(paragraph._p.pPr.find(qn("w:" + tag)).get(qn("w:val")), "0")
        order = [node.tag.split("}")[-1] for node in paragraph._p.pPr]
        self.assertLess(order.index("wordWrap"), order.index("autoSpaceDE"))
        self.assertLess(order.index("autoSpaceDN"), order.index("spacing"))
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "roundtrip.docx"
            document.save(path)
            self.assertEqual(docx.Document(path).tables[0].cell(0, 0).text, source)

    def test_headings_plain_prose_and_disabled_mode_keep_original_style(self):
        manager = StyleManager()
        for text, bold in (("教学内容 ls /etc/passwd", True), ("活动1 ls命令讲解", False), ("普通中文与Linux，24课时", False)):
            cell = docx.Document().add_table(rows=1, cols=1).cell(0, 0)
            manager.format_cell(cell, text, bold=bold)
            self.assertEqual(len([run for run in cell.paragraphs[0].runs if run.text]), 1)
            self.assertIsNone(cell.paragraphs[0]._p.pPr.find(qn("w:wordWrap")))
        disabled = StyleManager(style_override={"technical_text": {"enabled": False}})
        cell = docx.Document().add_table(rows=1, cols=1).cell(0, 0)
        disabled.format_cell(cell, "执行ls -l /etc/passwd")
        self.assertEqual(len([run for run in cell.paragraphs[0].runs if run.text]), 1)

    def test_schedule_body_and_rich_cell_items_support_technical_text(self):
        manager = StyleManager()
        cell = docx.Document().add_table(rows=1, cols=1).cell(0, 0)
        manager.format_prog_cell(cell, "查看/etc/passwd")
        self.assertEqual(cell.text, "查看/etc/passwd")
        self.assertEqual(tokens(cell.text), ["/etc/passwd"])
        manager.format_cell(cell, [{"text": "执行ls -l", "size_pt": 14}, ("访问/etc/hosts", False, "楷体", 13)])
        self.assertEqual(cell.text, "执行ls -l\n访问/etc/hosts")
        self.assertTrue(all(run.font.size.pt == 14 for run in cell.paragraphs[0].runs if run.text))
        self.assertTrue(all(run.font.size.pt == 13 for run in cell.paragraphs[1].runs if run.text))

    def test_font_name_xml_escaping_and_run_property_order(self):
        manager = StyleManager(style_override={"technical_text": {"font_name": 'Mono & "Code"'}})
        paragraph = docx.Document().add_paragraph()
        runs = manager.add_cell_text(paragraph, "ls", size_pt=12)
        fonts = runs[0]._r.rPr.find(qn("w:rFonts"))
        self.assertEqual(fonts.get(qn("w:ascii")), 'Mono & "Code"')
        order = [node.tag.split("}")[-1] for node in runs[0]._r.rPr]
        self.assertLess(order.index("rFonts"), order.index("b"))
        self.assertLess(order.index("color"), order.index("sz"))

    def test_technical_font_obeys_explicit_substitution_without_changing_cjk_font(self):
        manager = StyleManager(style_override={
            "technical_text": {"font_name": "Custom Code Font"},
            "fonts": {"substitutions": {"Custom Code Font": "Consolas"}},
        })
        paragraph = docx.Document().add_paragraph()
        manager.add_cell_text(paragraph, "执行ls -l", font_name="楷体", size_pt=14)
        self.assertEqual(paragraph.text, "执行ls -l")
        for run in paragraph.runs:
            self.assertEqual(run.font.size.pt, 14)
            fonts = run._r.rPr.find(qn("w:rFonts"))
            self.assertEqual(fonts.get(qn("w:eastAsia")), "楷体")
            if run.text in ("ls", "-l"):
                self.assertEqual(fonts.get(qn("w:ascii")), "Consolas")
                self.assertEqual(fonts.get(qn("w:hAnsi")), "Consolas")


if __name__ == "__main__":
    unittest.main()
